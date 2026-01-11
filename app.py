# app.py
# ------------------------------------------------------------
# AI People Reader - Presentation Analysis Report Generator
# Upload a video -> generate a DOCX report similar to your sample.
#
# Install:
#   pip install streamlit opencv-python python-docx numpy pandas matplotlib openpyxl
#
# Run:
#   streamlit run app.py
# ------------------------------------------------------------

# ===============================
# SAFE WORKER MODE ENTRY (TOP)
# ===============================
import os
import json
import uuid
import boto3
import streamlit as st
from datetime import datetime

def _require_env(keys):
    missing = [k for k in keys if not os.getenv(k)]
    if missing:
        st.warning("Worker mode ยังไม่พร้อม (ขาด env): " + ", ".join(missing))
        st.stop()

MODE = st.selectbox(
    "Select mode",
    ["Original app", "Worker mode (background processing)"],
    index=0
)

if MODE == "Worker mode (background processing)":
    st.title("🧠 AI People Reader — Worker Mode")

    _require_env(["AWS_ACCESS_KEY_ID", "AWS_SECRET_ACCESS_KEY", "AWS_REGION", "S3_BUCKET"])
    bucket = os.getenv("S3_BUCKET")

    def s3():
        return boto3.client(
            "s3",
            region_name=os.getenv("AWS_REGION"),
            aws_access_key_id=os.getenv("AWS_ACCESS_KEY_ID"),
            aws_secret_access_key=os.getenv("AWS_SECRET_ACCESS_KEY"),
        )

    video = st.file_uploader("Upload video for background processing", type=["mp4", "mov", "m4v"])

    if video and st.button("Submit to worker"):
        job_id = datetime.utcnow().strftime("%Y%m%d_%H%M%S_") + uuid.uuid4().hex[:6]
        input_key = f"jobs/{job_id}/input/input.mp4"

        s3().upload_fileobj(video, bucket, input_key)

        # initial status
        s3().put_object(
            Bucket=bucket,
            Key=f"jobs/{job_id}/status.json",
            Body=json.dumps({
                "job_id": job_id,
                "status": "queued",
                "progress": 0,
                "message": "Queued",
                "updated_at": datetime.utcnow().isoformat() + "Z",
                "outputs": {}
            }, ensure_ascii=False).encode("utf-8"),
            ContentType="application/json"
        )

        # enqueue ticket
        s3().put_object(
            Bucket=bucket,
            Key=f"jobs/pending/{job_id}.json",
            Body=json.dumps({
                "job_id": job_id,
                "input_key": input_key,
                "created_at": datetime.utcnow().isoformat() + "Z"
            }, ensure_ascii=False).encode("utf-8"),
            ContentType="application/json"
        )

        st.success(f"Job submitted: {job_id}")
        st.session_state["job_id"] = job_id

    job_id = st.session_state.get("job_id")
    if job_id:
        st.markdown("### Job status")
        status_key = f"jobs/{job_id}/status.json"
        try:
            obj = s3().get_object(Bucket=bucket, Key=status_key)
            status = json.loads(obj["Body"].read().decode("utf-8"))
            st.json(status)
            st.progress(int(status.get("progress", 0)))

            if status.get("status") == "done":
                overlay_key = status.get("outputs", {}).get("overlay_key")
                if overlay_key:
                    url = s3().generate_presigned_url(
                        "get_object",
                        Params={"Bucket": bucket, "Key": overlay_key},
                        ExpiresIn=3600
                    )
                    st.markdown(f"[⬇️ Download overlay]({url})")
        except Exception:
            st.info("Waiting for worker... (refresh this page)")

    st.stop()
# ===============================
# END WORKER MODE ENTRY
# ===============================

import os
import io
import math
import shutil
import tempfile
import base64
from dataclasses import dataclass, asdict
from datetime import datetime

import streamlit as st
import cv2
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK

try:
    import mediapipe as mp  # type: ignore
    # Some environments can accidentally install/import a different `mediapipe` module
    # that does not expose the public API we need (e.g., missing `solutions`).
    # In that case, treat MediaPipe as unavailable and fall back gracefully.
    if not hasattr(mp, "solutions"):
        mp = None
except Exception:
    mp = None

try:
    from moviepy.editor import VideoFileClip  # type: ignore
except Exception:
    VideoFileClip = None



# =========================
# Report content templates (to match Report.pdf sample wording)
# =========================
def _scale_label(scale: str) -> str:
    s = (scale or "").strip().lower()
    if s.startswith("high"):
        return "High"
    if s.startswith("mod"):
        return "Moderate"
    if s.startswith("low"):
        return "Low"
    # fallback for unexpected labels
    return (scale or "—").strip() or "—"


REPORT_CATEGORY_TEMPLATES: dict[str, dict[str, object]] = {
    # Keys are CategoryResult.name_en
    "Engaging & Connecting": {
        "bullets": [
            "Approachability",
            "Relatability",
            "Engagement, connect and build instant rapport with team",
        ],
    },
    "Confidence": {
        "bullets": [
            "Optimistic Presence",
            "Focus",
            "Ability to persuade and stand one’s ground, in order to convince others.",
        ],
    },
    "Authority": {
        "bullets": [
            "Showing sense of importance and urgency in subject matter",
            "Pressing for action",
        ],
    },
}


# =========================
# Data model
# =========================
@dataclass
class CategoryResult:
    name_en: str
    name_th: str
    score: int
    scale: str
    positives: int
    total: int
    description: str = ""


@dataclass
class ReportData:
    client_name: str
    analysis_date: str
    video_length_str: str
    overall_score: int
    categories: list
    summary_comment: str
    generated_by: str


# =========================
# Helpers
# =========================
def safe_int(x, default=0):
    try:
        return int(x)
    except Exception:
        return default


def format_seconds_to_mmss(total_seconds: float) -> str:
    total_seconds = max(0, float(total_seconds))
    mm = int(total_seconds // 60)
    ss = int(round(total_seconds - mm * 60))
    if ss == 60:
        mm += 1
        ss = 0
    return f"{mm:02d}:{ss:02d}"


def get_video_duration_seconds(video_path: str) -> float:
    cap = cv2.VideoCapture(video_path)
    if not cap.isOpened():
        return 0.0
    fps = cap.get(cv2.CAP_PROP_FPS) or 0.0
    frames = cap.get(cv2.CAP_PROP_FRAME_COUNT) or 0.0
    cap.release()
    if fps <= 0:
        return 0.0
    return float(frames / fps)


def generate_dot_motion_video_bytes(
    input_video_path: str,
    dot_size: int = 2,
    progress_cb=None,
) -> tuple[bytes, np.ndarray | None]:
    """
    Create a "dot motion" visualization video: black background + white dots at MediaPipe Pose landmarks.
    Returns (mp4_bytes, preview_frame_rgb_or_None).
    """
    if mp is None:
        raise RuntimeError("MediaPipe is not installed. Install `mediapipe` to enable dot motion generation.")

    mp_pose = mp.solutions.pose

    cap = cv2.VideoCapture(input_video_path)
    if not cap.isOpened():
        raise RuntimeError("Cannot open video for dot motion processing.")

    fps = int(cap.get(cv2.CAP_PROP_FPS) or 30)
    width = int(cap.get(cv2.CAP_PROP_FRAME_WIDTH) or 0)
    height = int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT) or 0)
    total_frames = int(cap.get(cv2.CAP_PROP_FRAME_COUNT) or 0)
    if width <= 0 or height <= 0:
        cap.release()
        raise RuntimeError("Could not read video dimensions.")

    # Write processed video (no audio) to temp
    no_audio_path = None
    final_path = None
    preview_frame = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".mp4") as tmp_out:
            no_audio_path = tmp_out.name

        fourcc = cv2.VideoWriter_fourcc(*"mp4v")
        out = cv2.VideoWriter(no_audio_path, fourcc, fps, (width, height))
        if not out.isOpened():
            raise RuntimeError("Could not create output video writer.")

        frames_processed = 0
        with mp_pose.Pose(static_image_mode=False) as pose:
            while cap.isOpened():
                ret, frame = cap.read()
                if not ret:
                    break
                h, w = frame.shape[:2]
                output = np.zeros((h, w, 3), dtype=np.uint8)

                rgb = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
                results = pose.process(rgb)

                if results.pose_landmarks:
                    for lm in results.pose_landmarks.landmark:
                        cx, cy = int(lm.x * w), int(lm.y * h)
                        if 0 <= cx < w and 0 <= cy < h:
                            cv2.circle(
                                output,
                                (cx, cy),
                                radius=int(dot_size),
                                color=(255, 255, 255),
                                thickness=-1,
                            )

                out.write(output)

                if frames_processed % 10 == 0:
                    preview_frame = cv2.cvtColor(output, cv2.COLOR_BGR2RGB)

                frames_processed += 1
                if progress_cb and total_frames > 0:
                    progress_cb(frames_processed / total_frames, frames_processed, total_frames)

        cap.release()
        out.release()

        # Try to add original audio (best-effort)
        if VideoFileClip is not None:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".mp4") as tmp_final:
                final_path = tmp_final.name

            original_video = VideoFileClip(input_video_path)
            processed_video = VideoFileClip(no_audio_path)
            try:
                final_video = processed_video.set_audio(original_video.audio) if original_video.audio is not None else processed_video
                final_video.write_videofile(
                    final_path,
                    codec="libx264",
                    audio_codec="aac",
                    temp_audiofile="temp-audio.m4a",
                    remove_temp=True,
                    verbose=False,
                    logger=None,
                )
            finally:
                try:
                    original_video.close()
                except Exception:
                    pass
                try:
                    processed_video.close()
                except Exception:
                    pass
                try:
                    final_video.close()
                except Exception:
                    pass

            out_path = final_path
        else:
            out_path = no_audio_path

        with open(out_path, "rb") as f:
            mp4_bytes = f.read()

        return mp4_bytes, preview_frame
    finally:
        try:
            cap.release()
        except Exception:
            pass
        # Cleanup temp files
        for p in [final_path, no_audio_path]:
            if p and os.path.exists(p):
                try:
                    os.unlink(p)
                except Exception:
                    pass
        if os.path.exists("temp-audio.m4a"):
            try:
                os.unlink("temp-audio.m4a")
            except Exception:
                pass


def recommend_sampling_settings(duration_seconds: float, preset: str = "Balanced") -> tuple[int, int]:
    """
    Recommend (sample_fps, max_frames) based on uploaded video duration and a preset.
    Goal: analyze the full clip with reasonable density, while keeping runtime manageable.

    Returns:
        sample_fps: int in [1..10]
        max_frames: int in [50..1500]
    """
    d = float(max(0.0, duration_seconds))
    p = (preset or "Balanced").strip().lower()

    # Presets tune how densely we sample and how much we cap max_frames.
    # - Fast: fewer samples, lower cap (good for batching many videos)
    # - Balanced: default
    # - Accurate: more samples, higher cap (slower)
    if p.startswith("fast"):
        max_cap = 600
        if d <= 45:
            sample_fps = 6
        elif d <= 90:
            sample_fps = 5
        elif d <= 180:
            sample_fps = 4
        elif d <= 300:
            sample_fps = 3
        else:
            sample_fps = 2
    elif p.startswith("accur"):
        max_cap = 1500
        if d <= 45:
            sample_fps = 10
        elif d <= 90:
            sample_fps = 10
        elif d <= 180:
            sample_fps = 8
        elif d <= 300:
            sample_fps = 6
        else:
            sample_fps = 4
    else:
        # Balanced
        max_cap = 1200
        if d <= 45:
            sample_fps = 10
        elif d <= 90:
            sample_fps = 8
        elif d <= 180:
            sample_fps = 6
        elif d <= 300:
            sample_fps = 4
        else:
            sample_fps = 3

    # Try to cover the full video at the selected sample rate
    max_frames = int(math.ceil(d * sample_fps)) if d > 0 else 300

    # Clamp to UI limits
    sample_fps = int(max(1, min(10, sample_fps)))
    max_frames = int(max(50, min(1500, min(max_cap, max_frames))))
    return sample_fps, max_frames


def recommend_advanced_settings(preset: str = "Balanced") -> dict[str, float | int]:
    """
    Recommend advanced accuracy settings by preset.
    These trade speed vs robustness vs accuracy.
    """
    p = (preset or "Balanced").strip().lower()
    if p.startswith("fast"):
        return {
            "pose_model_complexity": 0,
            "pose_min_detection_confidence": 0.5,
            "pose_min_tracking_confidence": 0.5,
            "face_min_detection_confidence": 0.5,
            "facemesh_min_detection_confidence": 0.5,
            "facemesh_min_tracking_confidence": 0.5,
        }
    if p.startswith("accur"):
        return {
            "pose_model_complexity": 2,
            "pose_min_detection_confidence": 0.7,
            "pose_min_tracking_confidence": 0.7,
            "face_min_detection_confidence": 0.7,
            "facemesh_min_detection_confidence": 0.7,
            "facemesh_min_tracking_confidence": 0.7,
        }
    # Balanced
    return {
        "pose_model_complexity": 1,
        "pose_min_detection_confidence": 0.6,
        "pose_min_tracking_confidence": 0.6,
        "face_min_detection_confidence": 0.6,
        "facemesh_min_detection_confidence": 0.6,
        "facemesh_min_tracking_confidence": 0.6,
    }


# =========================
# Excel Reference Data Loader
# =========================
def load_effort_reference(excel_path: str = None) -> pd.DataFrame:
    """Load Effort motion reference data from Excel file."""
    if excel_path is None:
        excel_path = os.path.join(os.path.dirname(__file__), "Effort.xlsx")
    
    if not os.path.exists(excel_path):
        raise FileNotFoundError(f"Effort.xlsx not found at {excel_path}")
    
    df = pd.read_excel(excel_path, header=None)
    # Extract header row (row 1) and data rows (rows 2+)
    df.columns = ['Motion Type', 'Direction', 'Body Part Involvement', 'Pathway', 'Timing', 'Other Motion Clues']
    df = df.iloc[2:].reset_index(drop=True)  # Skip header rows, keep data
    return df


def load_shape_reference(excel_path: str = None) -> pd.DataFrame:
    """Load Shape motion reference data from Excel file."""
    if excel_path is None:
        excel_path = os.path.join(os.path.dirname(__file__), "Shape.xlsx")
    
    if not os.path.exists(excel_path):
        raise FileNotFoundError(f"Shape.xlsx not found at {excel_path}")
    
    df = pd.read_excel(excel_path, header=None)
    # Extract header row (row 1) and data rows (rows 2+)
    df.columns = ['Motion Type', 'Direction', 'Body Part Involvement', 'Pathway', 'Timing', 'Other Motion Clues']
    df = df.iloc[2:].reset_index(drop=True)  # Skip header rows, keep data
    return df


# =========================
# Graph Generation Functions
# =========================
def generate_effort_graph(effort_detection: dict, shape_detection: dict, output_path: str):
    """
    Generate Graph 1.png showing combined Effort and Shape motion detection results.
    Two-panel layout: left shows all motions, right shows top 3.
    Matches exact format and proportions from the example image.
    
    Args:
        effort_detection: Dict with effort motion counts/percentages
        shape_detection: Dict with shape motion counts/percentages
        output_path: Path to save the graph
    """
    # Load reference data
    effort_df = load_effort_reference()
    shape_df = load_shape_reference()
    
    # Get effort motion types and filter out excluded ones
    effort_motions = effort_df['Motion Type'].tolist()
    excluded_motions = ['Floating', 'Slashing', 'Wringing']
    effort_motions = [m for m in effort_motions if m not in excluded_motions]
    
    # Get shape motion types
    shape_motions = shape_df['Motion Type'].tolist()
    
    # Combine all motion types
    all_motion_types = effort_motions + shape_motions
    
    # Get detection counts for all motions
    all_counts = []
    for motion in all_motion_types:
        count = 0
        # Check if it's an effort motion
        if motion in effort_motions:
            count = effort_detection.get(motion, 0)
            if count == 0:
                normalized_name = motion.lower().replace(' ', '_')
                key = f"{normalized_name}_count"
                count = effort_detection.get(key, 0)
        # Otherwise it's a shape motion
        elif motion in shape_motions:
            count = shape_detection.get(motion, 0)
            if count == 0:
                normalized_name = motion.lower().replace(' ', '_')
                key = f"{normalized_name}_count"
                count = shape_detection.get(key, 0)
        
        all_counts.append(count)
    
    # Calculate percentages
    total = sum(all_counts) if sum(all_counts) > 0 else 1  # Avoid division by zero
    percentages = [(count / total) * 100 for count in all_counts]
    
    # Sort by percentage (high to low)
    sorted_data = sorted(zip(all_motion_types, percentages, all_counts), key=lambda x: x[1], reverse=True)
    if sorted_data:
        sorted_motions, sorted_percentages, sorted_counts = zip(*sorted_data)
    else:
        sorted_motions, sorted_percentages, sorted_counts = all_motion_types, percentages, all_counts
    
    # Get top 3
    top3_motions = list(sorted_motions[:3])
    top3_percentages = list(sorted_percentages[:3])
    
    # Create figure with two subplots side by side - matching exact proportions
    num_all_motions = len(sorted_motions)
    # Use fixed proportions that match the example image
    fig_width = 14
    fig_height = max(7, num_all_motions * 0.45)  # Adjust for number of bars
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(fig_width, fig_height), 
                                    gridspec_kw={'width_ratios': [1.2, 0.8], 'wspace': 0.4})
    
    # Calculate bar height to match between panels - use same height for both
    bar_height = 0.7
    
    # Left panel: All motions (light blue bars)
    y_pos_all = range(len(sorted_motions))
    bars_all = ax1.barh(y_pos_all, sorted_percentages, color='#7EC8FF', height=bar_height)  # Light blue
    
    ax1.set_yticks(y_pos_all)
    ax1.set_yticklabels(sorted_motions, fontsize=11)
    ax1.set_xlabel('Percentage (%)', fontsize=12, fontweight='bold')
    ax1.set_xlim(0, 100)
    ax1.set_xticks([0, 20, 40, 60, 80, 100])
    ax1.set_ylim(-0.5, len(sorted_motions) - 0.5)  # Set y-axis limits
    ax1.set_title('Effort Summary', fontsize=14, fontweight='bold', pad=15)
    
    # Remove all spines for clean look
    ax1.spines['top'].set_visible(False)
    ax1.spines['right'].set_visible(False)
    ax1.spines['bottom'].set_visible(True)
    ax1.spines['left'].set_visible(True)
    ax1.spines['bottom'].set_color('#E0E0E0')
    ax1.spines['left'].set_color('#E0E0E0')
    
    # Light gray grid lines
    ax1.grid(axis='x', alpha=0.3, linestyle='-', linewidth=0.5, color='#E0E0E0')
    ax1.set_axisbelow(True)
    ax1.invert_yaxis()
    
    # Add percentage labels in green for all motions
    for bar, pct in zip(bars_all, sorted_percentages):
        if pct > 0:
            width = bar.get_width()
            ax1.text(width + 1, bar.get_y() + bar.get_height()/2.,
                    f'{pct:.1f}%',
                    ha='left', va='center', fontsize=11, fontweight='bold', color='#0B3D91')
    
    # Right panel: Top 3 (light blue bars) - use same y-positions and same y-axis limits
    top3_labels = [f"{motion} - Rank #{i+1}" for i, motion in enumerate(top3_motions)]
    # Use the same y-positions as the top 3 in the left panel (positions 0, 1, 2)
    y_pos_top3 = [0, 1, 2]
    bars_top3 = ax2.barh(y_pos_top3, top3_percentages, color='#7EC8FF', height=bar_height)  # Light blue - same height
    
    # Match the y-axis limits to align with left panel - this ensures same bar size
    ax2.set_ylim(-0.5, len(sorted_motions) - 0.5)  # Same y-axis limits as left panel
    ax2.set_yticks(y_pos_top3)
    ax2.set_yticklabels(top3_labels, fontsize=11)
    ax2.set_xlabel('Percentage (%)', fontsize=12, fontweight='bold')
    ax2.set_xlim(0, 100)
    ax2.set_xticks([0, 20, 40, 60, 80, 100])
    ax2.set_title('Top Movement Efforts', fontsize=14, fontweight='bold', pad=15)
    
    # Remove all spines for clean look
    ax2.spines['top'].set_visible(False)
    ax2.spines['right'].set_visible(False)
    ax2.spines['bottom'].set_visible(True)
    ax2.spines['left'].set_visible(True)
    ax2.spines['bottom'].set_color('#E0E0E0')
    ax2.spines['left'].set_color('#E0E0E0')
    
    # Light gray grid lines
    ax2.grid(axis='x', alpha=0.3, linestyle='-', linewidth=0.5, color='#E0E0E0')
    ax2.set_axisbelow(True)
    ax2.invert_yaxis()
    
    # Add percentage labels in green for top 3
    for bar, pct in zip(bars_top3, top3_percentages):
        if pct > 0:
            width = bar.get_width()
            ax2.text(width + 1, bar.get_y() + bar.get_height()/2.,
                    f'{pct:.1f}%',
                    ha='left', va='center', fontsize=11, fontweight='bold', color='#0B3D91')
    
    # Set tick parameters for cleaner look
    ax1.tick_params(axis='both', which='major', labelsize=10, colors='black')
    ax2.tick_params(axis='both', which='major', labelsize=10, colors='black')
    
    # Tight layout for proper spacing
    plt.tight_layout()
    
    # Save with exact DPI and format matching reference
    plt.savefig(output_path, dpi=300, bbox_inches='tight', facecolor='white', edgecolor='none')
    plt.close()


def generate_shape_graph(detection_results: dict, output_path: str):
    """
    Generate Graph 2.png showing Shape motion detection results.
    Matches the exact format of the reference Graph 2.png.
    
    Args:
        detection_results: Dict with shape motion counts/percentages
        output_path: Path to save the graph
    """
    shape_df = load_shape_reference()
    motion_types = shape_df['Motion Type'].tolist()
    
    # Get detection counts from results
    counts = []
    for motion in motion_types:
        # Try to get count using motion name directly, or normalized key
        count = detection_results.get(motion, 0)
        if count == 0:
            normalized_name = motion.lower().replace(' ', '_')
            key = f"{normalized_name}_count"
            count = detection_results.get(key, 0)
        counts.append(count)
    
    # Calculate percentages
    total = sum(counts) if sum(counts) > 0 else 1  # Avoid division by zero
    percentages = [(count / total) * 100 for count in counts]
    
    # Sort by percentage (high to low) and maintain corresponding motion types and colors
    sorted_data = sorted(zip(motion_types, percentages, counts), key=lambda x: x[1], reverse=True)
    if sorted_data:
        sorted_motions, sorted_percentages, sorted_counts = zip(*sorted_data)
    else:
        sorted_motions, sorted_percentages, sorted_counts = motion_types, percentages, counts
    
    # Create bar chart with standardized format (same as Graph 1)
    num_motions = len(sorted_motions)
    # Scale width: ~1.5 inches per motion, minimum 12 inches
    fig_width = max(12, num_motions * 1.5)
    # Maintain good aspect ratio for readability
    fig_height = 7.14
    fig, ax = plt.subplots(figsize=(fig_width, fig_height))
    
    # Use a consistent light blue across the whole graph (per user request)
    colors = ['#7EC8FF' for _ in sorted_motions]
    
    # Adjust bar width based on number of motions (narrower for more bars)
    bar_width = 0.5 if num_motions > 4 else 0.6
    bars = ax.bar(sorted_motions, sorted_percentages, color=colors, width=bar_width)
    
    # Set labels and title with exact formatting
    ax.set_xlabel('Shape Motion Type', fontsize=12, fontweight='bold')
    ax.set_ylabel('Percentage (%)', fontsize=12, fontweight='bold')
    ax.set_title('Shape Motion Detection Results', fontsize=14, fontweight='bold', pad=20)
    
    # Grid styling to match reference
    ax.grid(axis='y', alpha=0.3, linestyle='-', linewidth=0.5)
    ax.set_axisbelow(True)
    
    # Remove top and right spines for cleaner look
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    
    # Add percentage labels on top of bars
    for bar in bars:
        height = bar.get_height()
        if height > 0:
            ax.text(bar.get_x() + bar.get_width()/2., height,
                    f'{height:.1f}%',
                    ha='center', va='bottom', fontsize=11, fontweight='bold', color='#0B3D91')
    
    # X-axis formatting - rotate if many motion types to avoid overlap
    if num_motions > 6:
        plt.xticks(rotation=15, ha='right', fontsize=10)
    else:
        plt.xticks(rotation=0, fontsize=11)
    plt.yticks(fontsize=10)
    
    # Ensure y-axis starts from 0 and goes up to 100% (or slightly above max)
    max_percentage = max(sorted_percentages) if sorted_percentages else 100
    ax.set_ylim(bottom=0, top=min(100, max_percentage * 1.15))
    
    # Tight layout for proper spacing
    plt.tight_layout()
    
    # Save with exact DPI and format matching reference
    plt.savefig(output_path, dpi=300, bbox_inches='tight', facecolor='white', edgecolor='none')
    plt.close()


# =========================
# Placeholder analysis engine
# (Replace this with your real detection logic)
# =========================

# SCORING CRITERIA DOCUMENTATION:
# ===============================
# Each category is scored from 1-10 based on detected positive indicators.
# 
# 1. ENGAGING & CONNECTING (การสร้างความเป็นมิตรและสร้างสัมพันธภาพ)
#    Criteria: Positive indicators include:
#    - Open body posture (arms not crossed, open chest)
#    - Forward lean toward audience
#    - Smiling or positive facial expressions
#    - Gestures that include/invite (open palms, arms spread)
#    - Eye contact with audience
#    - Nodding in agreement/acknowledgment
#    - Approachable stance (body oriented toward audience)
#    Score Calculation: (positive_indicators / total_indicators) * 10
#    Scale: 1-3 = low, 4-6 = moderate, 7-10 = high
#
# 2. INFORMATION SHARING (การให้ข้อมูล)
#    Criteria: Positive indicators include:
#    - Clear pointing gestures toward visual aids
#    - Descriptive hand movements (showing size, shape, direction)
#    - Organized movement patterns (methodical, structured)
#    - Stable stance during information delivery
#    - Controlled gestures that emphasize key points
#    - Pauses for emphasis
#    Score Calculation: (positive_indicators / total_indicators) * 10
#    Scale: 1-3 = low, 4-6 = moderate, 7-10 = high
#
# 3. CONVINCING WITH CONFIDENCE (การโน้มน้าว)
#    Criteria: Positive indicators include:
#    - Confident, decisive movements
#    - Strong, grounded stance
#    - Assertive gestures (palm down, pointing)
#    - Forward momentum in movement
#    - Strong eye contact
#    - Upright, confident posture
#    - Controlled but impactful gestures
#    - Minimal hesitation or uncertainty in movement
#    Score Calculation: (positive_indicators / total_indicators) * 10
#    Scale: 1-3 = low, 4-6 = moderate, 7-10 = high
#
# 4. INSTRUCTING WITH AUTHORITY AND LEADERSHIP (การสั่ง)
#    Criteria: Positive indicators include:
#    - Commanding presence (strong, stable posture)
#    - Directive gestures (pointing, palm down, chopping motions)
#    - Dominant spatial positioning
#    - Minimal unnecessary movement (authoritative stillness)
#    - Strong, upright alignment
#    - Confident, deliberate movements
#    - Leadership-oriented stance (taking space, projecting authority)
#    Score Calculation: (positive_indicators / total_indicators) * 10
#    Scale: 1-3 = low, 4-6 = moderate, 7-10 = high
#
# TOTAL INDICATORS:
#    Should be calculated from actual video analysis:
#    - Frame count analysis (checking each frame for movement indicators)
#    - Keyframe sampling (analyzing representative frames)
#    - Temporal segmentation (analyzing movement patterns over time)
#    Currently: Placeholder using duration * 300 (frames/second estimation)
#
# OVERALL SCORE:
#    Average of all four category scores, rounded to nearest integer.
# ===============================

def analyze_video_motion_features(video_path: str) -> dict:
    """
    Analyze actual video features to detect motion patterns.
    Returns features like motion variance, movement intensity, etc.
    """
    cap = cv2.VideoCapture(video_path)
    if not cap.isOpened():
        return {"motion_variance": 0.0, "avg_brightness": 0.5, "movement_intensity": 0.0, "frame_count": 0}
    
    fps = cap.get(cv2.CAP_PROP_FPS) or 30.0
    frame_count = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
    
    # Sample frames for analysis (every Nth frame to speed up)
    sample_interval = max(1, int(fps / 5))  # ~5 frames per second
    frames_to_analyze = min(100, frame_count // sample_interval)  # Max 100 samples
    
    prev_frame = None
    motion_values = []
    brightness_values = []
    
    frame_idx = 0
    analyzed_count = 0
    
    while analyzed_count < frames_to_analyze and cap.isOpened():
        ret, frame = cap.read()
        if not ret:
            break
        
        if frame_idx % sample_interval == 0:
            # Convert to grayscale for motion analysis
            gray = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)
            
            # Calculate brightness
            brightness = np.mean(gray) / 255.0
            brightness_values.append(brightness)
            
            # Calculate motion if we have previous frame
            if prev_frame is not None:
                # Calculate frame difference
                diff = cv2.absdiff(gray, prev_frame)
                motion = np.mean(diff) / 255.0
                motion_values.append(motion)
            
            prev_frame = gray.copy()
            analyzed_count += 1
        
        frame_idx += 1
    
    cap.release()
    
    # Calculate statistics
    motion_variance = float(np.var(motion_values)) if motion_values else 0.0
    avg_brightness = float(np.mean(brightness_values)) if brightness_values else 0.5
    movement_intensity = float(np.mean(motion_values)) if motion_values else 0.0
    
    return {
        "motion_variance": motion_variance,
        "avg_brightness": avg_brightness,
        "movement_intensity": movement_intensity,
        "frame_count": frame_count,
        "analyzed_frames": analyzed_count,
    }


def _clamp(v: float, lo: float, hi: float) -> float:
    return float(max(lo, min(hi, v)))


def _safe_div(a: float, b: float, default: float = 0.0) -> float:
    return float(a / b) if b else float(default)


def _rotation_matrix_to_euler_degrees(R: np.ndarray) -> tuple[float, float, float]:
    """
    Convert rotation matrix to Euler angles (pitch, yaw, roll) in degrees.
    Uses a common convention (x=pitch, y=yaw, z=roll).
    """
    # Guard against numerical issues
    sy = math.sqrt(float(R[0, 0] * R[0, 0] + R[1, 0] * R[1, 0]))
    singular = sy < 1e-6
    if not singular:
        x = math.atan2(float(R[2, 1]), float(R[2, 2]))
        y = math.atan2(float(-R[2, 0]), float(sy))
        z = math.atan2(float(R[1, 0]), float(R[0, 0]))
    else:
        x = math.atan2(float(-R[1, 2]), float(R[1, 1]))
        y = math.atan2(float(-R[2, 0]), float(sy))
        z = 0.0
    return (math.degrees(x), math.degrees(y), math.degrees(z))


def _estimate_head_pose_from_facemesh(
    face_landmarks,
    image_w: int,
    image_h: int,
) -> tuple[float, float, float] | None:
    """
    Estimate head pose (pitch, yaw, roll) from MediaPipe FaceMesh landmarks using solvePnP.
    Returns degrees or None if estimation fails.
    """
    # Indices commonly used for head pose with FaceMesh
    # nose tip, chin, left eye outer, right eye outer, left mouth, right mouth
    idxs = [1, 152, 33, 263, 61, 291]
    try:
        pts_2d = []
        for idx in idxs:
            lm = face_landmarks.landmark[idx]
            pts_2d.append([lm.x * image_w, lm.y * image_h])
        image_points = np.array(pts_2d, dtype=np.float64)
    except Exception:
        return None

    # Approximate 3D model points (generic face model; units arbitrary but consistent)
    model_points = np.array(
        [
            (0.0, 0.0, 0.0),        # Nose tip
            (0.0, -63.6, -12.5),    # Chin
            (-43.3, 32.7, -26.0),   # Left eye outer corner
            (43.3, 32.7, -26.0),    # Right eye outer corner
            (-28.9, -28.9, -24.1),  # Left mouth corner
            (28.9, -28.9, -24.1),   # Right mouth corner
        ],
        dtype=np.float64,
    )

    # Camera matrix approximation
    focal_length = float(image_w)
    center = (float(image_w) / 2.0, float(image_h) / 2.0)
    camera_matrix = np.array(
        [[focal_length, 0, center[0]], [0, focal_length, center[1]], [0, 0, 1]],
        dtype=np.float64,
    )
    dist_coeffs = np.zeros((4, 1), dtype=np.float64)

    ok, rvec, tvec = cv2.solvePnP(
        model_points,
        image_points,
        camera_matrix,
        dist_coeffs,
        flags=cv2.SOLVEPNP_ITERATIVE,
    )
    if not ok:
        return None

    R, _ = cv2.Rodrigues(rvec)
    pitch, yaw, roll = _rotation_matrix_to_euler_degrees(R)
    return (float(pitch), float(yaw), float(roll))


def analyze_video_mediapipe(
    video_path: str,
    sample_fps: float = 5.0,
    max_frames: int = 300,
    pose_model_complexity: int = 1,
    pose_min_detection_confidence: float = 0.5,
    pose_min_tracking_confidence: float = 0.5,
    face_min_detection_confidence: float = 0.5,
    facemesh_min_detection_confidence: float = 0.5,
    facemesh_min_tracking_confidence: float = 0.5,
) -> dict:
    """
    Real analysis using MediaPipe Pose + FaceMesh.
    Produces:
      - category positives/total indicators and scores
      - presence insights from measured head pose/pose stability
      - graph detection dictionaries (effort/shape) derived from pose motion
    """
    if mp is None:
        raise RuntimeError("MediaPipe is not installed. Install mediapipe to enable real analysis.")

    cap = cv2.VideoCapture(video_path)
    if not cap.isOpened():
        raise RuntimeError("Cannot open uploaded video.")

    fps = float(cap.get(cv2.CAP_PROP_FPS) or 30.0)
    frame_count = int(cap.get(cv2.CAP_PROP_FRAME_COUNT) or 0)
    duration = float(frame_count / fps) if fps > 0 else get_video_duration_seconds(video_path)

    step = max(1, int(round(_safe_div(fps, sample_fps, default=6.0))))
    mp_pose = mp.solutions.pose
    mp_face = mp.solutions.face_mesh
    mp_face_det = mp.solutions.face_detection

    # Tracking buffers
    analyzed = 0
    pose_ok_frames = 0
    face_detected_frames = 0
    head_pose_ok_frames = 0
    eye_contact_frames = 0
    eye_contact_frames_headpose = 0
    eye_contact_frames_posefallback = 0

    hip_xs: list[float] = []
    shoulder_widths: list[float] = []
    arm_spans: list[float] = []
    yaws: list[float] = []
    pitches: list[float] = []

    # Motion for effort/shape
    prev_wrist: dict[str, tuple[float, float] | None] = {"L": None, "R": None}
    prev_wrist_v: dict[str, tuple[float, float] | None] = {"L": None, "R": None}
    prev_shoulder_width: float | None = None
    prev_arm_span: float | None = None
    prev_hip: tuple[float, float] | None = None
    prev_hip_v: tuple[float, float] | None = None

    effort_counts = {"Gliding": 0, "Punching": 0, "Dabbing": 0, "Flicking": 0, "Pressing": 0}
    shape_counts = {"Advancing": 0, "Retreating": 0, "Enclosing": 0, "Spreading": 0, "Directing": 0, "Indirecting": 0}

    # Per-second time series (one row per second)
    # Structure: {sec: {"Gliding": n, ..., "Indirecting": n}}
    per_second: dict[int, dict[str, int]] = {}
    current_sec = 0

    def _ensure_sec(sec: int):
        if sec not in per_second:
            per_second[sec] = {**{k: 0 for k in effort_counts.keys()}, **{k: 0 for k in shape_counts.keys()}}

    def _inc_effort(motion: str, n: int = 1):
        if motion in effort_counts:
            effort_counts[motion] += int(n)
            _ensure_sec(current_sec)
            per_second[current_sec][motion] += int(n)

    def _inc_shape(motion: str, n: int = 1):
        if motion in shape_counts:
            shape_counts[motion] += int(n)
            _ensure_sec(current_sec)
            per_second[current_sec][motion] += int(n)

    # Category indicator counts
    engaging_pos = 0
    info_pos = 0
    convince_pos = 0
    authority_pos = 0

    # Per-frame indicator totals (fixed)
    ENGAGING_IND = 3  # eye contact, open posture, gesture presence
    # Adaptability: (1) movement variety (effort/shape types), (2) focus (stable gaze), (3) guidance (directing / structured gestures)
    INFO_IND = 3
    CONVINCE_IND = 3  # eye contact, strong gesture, upright
    AUTH_IND = 3      # grounded stance, upright, calm gestures

    with mp_pose.Pose(
        static_image_mode=False,
        model_complexity=int(pose_model_complexity),
        enable_segmentation=False,
        min_detection_confidence=float(pose_min_detection_confidence),
        min_tracking_confidence=float(pose_min_tracking_confidence),
    ) as pose, mp_face_det.FaceDetection(
        model_selection=0,
        min_detection_confidence=float(face_min_detection_confidence),
    ) as face_det, mp_face.FaceMesh(
        static_image_mode=False,
        max_num_faces=1,
        refine_landmarks=True,
        min_detection_confidence=float(facemesh_min_detection_confidence),
        min_tracking_confidence=float(facemesh_min_tracking_confidence),
    ) as face_mesh:
        idx = 0
        while cap.isOpened() and analyzed < max_frames:
            ret, frame = cap.read()
            if not ret:
                break
            if idx % step != 0:
                idx += 1
                continue

            image_h, image_w = frame.shape[:2]
            rgb = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
            # Timestamp in seconds for this sampled frame
            pos_msec = float(cap.get(cv2.CAP_PROP_POS_MSEC) or 0.0)
            current_sec = int(pos_msec // 1000.0)
            _ensure_sec(current_sec)

            pose_res = pose.process(rgb)

            # --- Face / head pose ---
            eye_contact = False
            stable_gaze = False
            head_pose = None
            eye_contact_counted = False

            # Robust face detection -> crop -> FaceMesh -> head pose
            face_det_res = face_det.process(rgb)
            if face_det_res.detections:
                det = max(face_det_res.detections, key=lambda d: float(d.score[0]) if d.score else 0.0)
                score = float(det.score[0]) if det.score else 0.0
                if score >= 0.5:
                    face_detected_frames += 1
                    bb = det.location_data.relative_bounding_box
                    x1 = int(bb.xmin * image_w)
                    y1 = int(bb.ymin * image_h)
                    w = int(bb.width * image_w)
                    h = int(bb.height * image_h)
                    # Add margin around the face
                    mx = int(0.25 * w)
                    my = int(0.25 * h)
                    x1 = max(0, x1 - mx)
                    y1 = max(0, y1 - my)
                    x2 = min(image_w, x1 + w + 2 * mx)
                    y2 = min(image_h, y1 + h + 2 * my)
                    face_crop = rgb[y1:y2, x1:x2]
                    if face_crop.size > 0:
                        face_res = face_mesh.process(face_crop)
                        if face_res.multi_face_landmarks:
                            head_pose = _estimate_head_pose_from_facemesh(
                                face_res.multi_face_landmarks[0],
                                face_crop.shape[1],
                                face_crop.shape[0],
                            )
                            if head_pose is not None:
                                pitch, yaw, roll = head_pose
                                pitches.append(pitch)
                                yaws.append(yaw)
                                # "Eye contact" approximation: looking near camera
                                eye_contact = (abs(yaw) <= 20.0 and abs(pitch) <= 20.0)
                                head_pose_ok_frames += 1
                                if eye_contact:
                                    eye_contact_frames += 1
                                    eye_contact_frames_headpose += 1
                                    eye_contact_counted = True

                                # gaze stability from last few yaws
                                if len(yaws) >= 5:
                                    recent = np.array(yaws[-5:], dtype=np.float32)
                                    stable_gaze = float(np.std(recent)) < 5.0

            # --- Pose / posture ---
            open_posture = False
            gesture_present = False
            strong_gesture = False
            controlled_gesture = False
            grounded_stance = False
            upright = False
            calm_gestures = False
            directing_move = False
            indirecting_move = False

            if pose_res.pose_landmarks:
                lm = pose_res.pose_landmarks.landmark
                pose_ok_frames += 1

                def xy(i: int) -> tuple[float, float]:
                    return (float(lm[i].x), float(lm[i].y))

                def vis(i: int) -> float:
                    return float(getattr(lm[i], "visibility", 0.0))

                # Indices (MediaPipe Pose)
                L_SH, R_SH = 11, 12
                L_EL, R_EL = 13, 14
                L_WR, R_WR = 15, 16
                L_HIP, R_HIP = 23, 24
                L_ANK, R_ANK = 27, 28

                l_sh, r_sh = xy(L_SH), xy(R_SH)
                l_wr, r_wr = xy(L_WR), xy(R_WR)
                l_hip, r_hip = xy(L_HIP), xy(R_HIP)
                l_ank, r_ank = xy(L_ANK), xy(R_ANK)

                sh_width = float(np.linalg.norm(np.array(l_sh) - np.array(r_sh)))
                sh_width = max(sh_width, 1e-6)
                shoulder_widths.append(sh_width)

                hip = ((l_hip[0] + r_hip[0]) / 2.0, (l_hip[1] + r_hip[1]) / 2.0)
                hip_xs.append(hip[0])

                # Arm span (wrist-to-wrist)
                arm_span = float(np.linalg.norm(np.array(l_wr) - np.array(r_wr)))
                arm_spans.append(arm_span)

                # Gesture presence if wrists visible
                gesture_present = (vis(L_WR) > 0.5 and vis(R_WR) > 0.5)

                # Open posture if wrists far apart relative to shoulders
                open_posture = (arm_span / sh_width) >= 1.25

                # Uprightness from shoulder-to-hip vector angle to vertical axis in image plane
                shoulder_center = ((l_sh[0] + r_sh[0]) / 2.0, (l_sh[1] + r_sh[1]) / 2.0)
                v = np.array([hip[0] - shoulder_center[0], hip[1] - shoulder_center[1]], dtype=np.float32)
                # vertical reference in image is (0, 1)
                v_norm = float(np.linalg.norm(v))
                if v_norm > 1e-6:
                    cosang = _clamp(float(v[1] / v_norm), -1.0, 1.0)
                    angle = math.degrees(math.acos(cosang))
                    upright = angle <= 15.0

                # Grounded stance from ankle distance relative to shoulder width
                foot_span = float(np.linalg.norm(np.array(l_ank) - np.array(r_ank)))
                grounded_stance = (foot_span / sh_width) >= 0.9

                # Wrist velocity / acceleration for gesture intensity (normalized)
                def update_wrist(side: str, cur: tuple[float, float]):
                    nonlocal strong_gesture, controlled_gesture, calm_gestures
                    prev = prev_wrist[side]
                    if prev is None:
                        prev_wrist[side] = cur
                        prev_wrist_v[side] = None
                        return
                    vx, vy = (cur[0] - prev[0]), (cur[1] - prev[1])
                    speed = float(math.sqrt(vx * vx + vy * vy)) / sh_width
                    vprev = prev_wrist_v[side]
                    accel = 0.0
                    if vprev is not None:
                        ax, ay = (vx - vprev[0]), (vy - vprev[1])
                        accel = float(math.sqrt(ax * ax + ay * ay)) / sh_width

                    # Heuristics:
                    # - strong gesture: fast movement
                    # - controlled gesture: moderate speed
                    # - calm gestures: low speed
                    if speed >= 0.08:
                        strong_gesture = True
                    if 0.03 <= speed < 0.08:
                        controlled_gesture = True
                    if speed < 0.03:
                        calm_gestures = True

                    # Effort detection heuristics
                    # Punching: high speed + high accel
                    if speed >= 0.09 and accel >= 0.05:
                        _inc_effort("Punching")
                    # Flicking: high speed + sharp direction change
                    if vprev is not None:
                        dot = vx * vprev[0] + vy * vprev[1]
                        n1 = math.sqrt(vx * vx + vy * vy) + 1e-9
                        n2 = math.sqrt(vprev[0] * vprev[0] + vprev[1] * vprev[1]) + 1e-9
                        ang = math.degrees(math.acos(_clamp(dot / (n1 * n2), -1.0, 1.0)))
                        if speed >= 0.07 and ang >= 70:
                            _inc_effort("Flicking")
                    # Dabbing: quick but small/compact gesture (moderate speed, high accel)
                    if 0.05 <= speed < 0.09 and accel >= 0.05:
                        _inc_effort("Dabbing")
                    # Gliding: sustained smooth motion (moderate speed, low accel)
                    if 0.03 <= speed < 0.07 and accel < 0.03:
                        _inc_effort("Gliding")
                    # Pressing: downward motion (in image coordinates, +y) with moderate speed
                    if vy > 0.0 and 0.03 <= speed < 0.08:
                        _inc_effort("Pressing")

                    prev_wrist[side] = cur
                    prev_wrist_v[side] = (vx, vy)

                update_wrist("L", l_wr)
                update_wrist("R", r_wr)

                # Shape detection heuristics based on:
                # - Advancing/Retreating: change in apparent scale (shoulder width)
                # - Spreading/Enclosing: change in arm span
                # - Directing/Indirecting: hip movement direction stability
                if prev_shoulder_width is not None:
                    dscale = (sh_width - prev_shoulder_width) / prev_shoulder_width
                    if dscale >= 0.02:
                        _inc_shape("Advancing")
                    elif dscale <= -0.02:
                        _inc_shape("Retreating")
                prev_shoulder_width = sh_width

                if prev_arm_span is not None:
                    dspan = (arm_span - prev_arm_span) / max(prev_arm_span, 1e-6)
                    if dspan >= 0.03:
                        _inc_shape("Spreading")
                    elif dspan <= -0.03:
                        _inc_shape("Enclosing")
                prev_arm_span = arm_span

                if prev_hip is not None:
                    hvx, hvy = (hip[0] - prev_hip[0]), (hip[1] - prev_hip[1])
                    hv = (hvx, hvy)
                    if prev_hip_v is not None:
                        dot = hvx * prev_hip_v[0] + hvy * prev_hip_v[1]
                        n1 = math.sqrt(hvx * hvx + hvy * hvy) + 1e-9
                        n2 = math.sqrt(prev_hip_v[0] * prev_hip_v[0] + prev_hip_v[1] * prev_hip_v[1]) + 1e-9
                        ang = math.degrees(math.acos(_clamp(dot / (n1 * n2), -1.0, 1.0)))
                        if ang < 35.0:
                            _inc_shape("Directing")
                            directing_move = True
                        elif ang >= 70.0:
                            _inc_shape("Indirecting")
                            indirecting_move = True
                    prev_hip_v = hv
                prev_hip = hip

                # Pose-only fallback eye contact (when head pose isn't available):
                # Use face landmarks in Pose (nose/eyes/ears) to estimate if the head is facing camera.
                if head_pose is None:
                    # Indices for Pose "face" points
                    NOSE = 0
                    L_EAR, R_EAR = 7, 8
                    L_EYE, R_EYE = 2, 5

                    if vis(NOSE) > 0.5 and vis(L_EAR) > 0.5 and vis(R_EAR) > 0.5 and vis(L_EYE) > 0.5 and vis(R_EYE) > 0.5:
                        nose = xy(NOSE)
                        l_ear, r_ear = xy(L_EAR), xy(R_EAR)
                        ear_dist = float(np.linalg.norm(np.array(l_ear) - np.array(r_ear)))
                        if ear_dist > 1e-6:
                            # If nose is centered between ears (symmetric), assume facing camera.
                            nose_to_l = float(np.linalg.norm(np.array(nose) - np.array(l_ear)))
                            nose_to_r = float(np.linalg.norm(np.array(nose) - np.array(r_ear)))
                            symmetry = abs(nose_to_l - nose_to_r) / ear_dist  # 0 is perfect
                            if symmetry <= 0.15:
                                eye_contact = True
                                if not eye_contact_counted:
                                    eye_contact_frames += 1
                                    eye_contact_frames_posefallback += 1
                                    eye_contact_counted = True

            # --- Category indicators (counted per analyzed frame) ---
            # Adaptability indicators based on your definition:
            # - ability to use different types of movement to engage listeners -> "movement variety"
            # - ability to stay focused -> "stable gaze"
            # - guide listeners regarding subject matter -> "guiding" (directing motion and/or structured gestures with stable gaze)
            sec_counts = per_second.get(current_sec, {})
            movement_types_used = sum(1 for v in sec_counts.values() if int(v) > 0)
            movement_variety = movement_types_used >= 2  # at least 2 distinct motion types within the current second
            guiding = bool(directing_move) or bool(controlled_gesture and stable_gaze)

            engaging_pos += int(eye_contact) + int(open_posture) + int(gesture_present)
            info_pos += int(movement_variety) + int(stable_gaze) + int(guiding)
            convince_pos += int(eye_contact) + int(strong_gesture) + int(upright)
            authority_pos += int(grounded_stance) + int(upright) + int(calm_gestures)

            analyzed += 1
            idx += 1

    cap.release()

    if analyzed == 0 or pose_ok_frames < max(3, int(0.2 * analyzed)):
        raise RuntimeError("Pose detection did not find a person reliably in this video.")

    # Global stability metrics
    sway_std = float(np.std(np.array(hip_xs, dtype=np.float32))) if hip_xs else 0.0
    yaw_std = float(np.std(np.array(yaws, dtype=np.float32))) if yaws else 0.0
    # Note: eye contact is only measurable on frames where head-pose estimation succeeded.
    eye_contact_pct_of_head_pose_frames = _safe_div(float(eye_contact_frames_headpose), float(max(1, head_pose_ok_frames)), 0.0) * 100.0
    eye_contact_pct_of_analyzed_frames = _safe_div(float(eye_contact_frames), float(max(1, analyzed)), 0.0) * 100.0

    def pos_to_score(pos: int, total: int) -> int:
        ratio = _safe_div(float(pos), float(max(1, total)), 0.0)
        return int(np.clip(round(ratio * 10), 1, 10))

    total_engaging = analyzed * ENGAGING_IND
    total_info = analyzed * INFO_IND
    total_convince = analyzed * CONVINCE_IND
    total_authority = analyzed * AUTH_IND

    engaging_score = pos_to_score(engaging_pos, total_engaging)
    info_score = pos_to_score(info_pos, total_info)
    convince_score = pos_to_score(convince_pos, total_convince)
    authority_score = pos_to_score(authority_pos, total_authority)
    overall_score = int(round(np.mean([engaging_score, info_score, convince_score, authority_score])))

    # Presence summary (measured)
    if head_pose_ok_frames == 0:
        gaze_stability = "Not available (head pose not estimated reliably)"
    else:
        gaze_stability = "Stable" if yaw_std < 6.0 else ("Moderately stable" if yaw_std < 10.0 else "Unstable (frequent shifts)")
    body_sway = "Very stable posture" if sway_std < 0.01 else ("Minimal sway" if sway_std < 0.02 else ("Moderate sway" if sway_std < 0.035 else "Excessive sway detected"))

    presence_comment = (
        f"Face detected (of analyzed frames): {_safe_div(float(face_detected_frames), float(max(1, analyzed)), 0.0) * 100.0:.0f}%\n"
        f"Head pose estimated (of analyzed frames): {_safe_div(float(head_pose_ok_frames), float(max(1, analyzed)), 0.0) * 100.0:.0f}%\n"
        f"Eye contact (of head-pose frames): {eye_contact_pct_of_head_pose_frames:.0f}%\n"
        f"Eye contact (of analyzed frames): {eye_contact_pct_of_analyzed_frames:.0f}%\n"
        f"Gaze stability: {gaze_stability}\n"
        f"Body sway: {body_sway}\n"
        f"Notes: This analysis is computed from detected face/pose landmarks on sampled frames (~{sample_fps:.0f} fps, {analyzed} frames). "
        f"Eye contact is computed from head pose when available, otherwise a pose-based fallback is used."
    )

    # Prepare graph dictionaries (counts only; graph code derives percentages)
    effort_detection = {}
    for k, v in effort_counts.items():
        effort_detection[k] = int(v)
        effort_detection[f"{k.lower()}_count"] = int(v)

    shape_detection = {}
    for k, v in shape_counts.items():
        shape_detection[k] = int(v)
        shape_detection[f"{k.lower()}_count"] = int(v)

    return {
        "duration_seconds": duration,
        "analyzed_frames": analyzed,
        "pose_ok_frames": pose_ok_frames,
        "face_detected_frames": face_detected_frames,
        "head_pose_ok_frames": head_pose_ok_frames,
        "eye_contact_frames": eye_contact_frames,
        "eye_contact_frames_headpose": eye_contact_frames_headpose,
        "eye_contact_frames_posefallback": eye_contact_frames_posefallback,
        "total_indicators": int(max(total_engaging, total_info, total_convince, total_authority)),
        "engaging_pos": int(engaging_pos),
        "info_pos": int(info_pos),
        "convince_pos": int(convince_pos),
        "authority_pos": int(authority_pos),
        "engaging_score": int(engaging_score),
        "info_score": int(info_score),
        "convince_score": int(convince_score),
        "authority_score": int(authority_score),
        "overall_score": int(overall_score),
        "presence_comment": presence_comment,
        "effort_detection": effort_detection,
        "shape_detection": shape_detection,
        "per_second_detection": per_second,
        "analysis_engine": "mediapipe_pose_facemesh",
    }


def analyze_video_placeholder(video_path: str, seed: int = 42) -> dict:
    """
    Returns a dict with:
      - total_indicators
      - positives per category
      - plus "presence" insights (eye contact, stance, uprightness)
    
    This function now uses actual video analysis to generate unique scores per video.
    For full implementation, integrate skeleton detection (e.g., MediaPipe, OpenPose)
    and Laban Movement Analysis parameters.
    """
    # Create unique seed based on video file properties
    import hashlib
    video_hash = hashlib.md5((video_path + str(os.path.getsize(video_path))).encode()).hexdigest()
    unique_seed = int(video_hash[:8], 16) % (2**31)
    rng = np.random.default_rng(unique_seed)

    # Analyze actual video features
    video_features = analyze_video_motion_features(video_path)
    
    duration = get_video_duration_seconds(video_path)
    
    # Use actual video features to generate realistic scores
    # Higher motion intensity suggests more engaging movement
    motion_factor = min(1.0, video_features["movement_intensity"] * 2.0)  # Normalize to 0-1
    
    # Motion variance suggests variety in movements (good for different categories)
    variance_factor = min(1.0, video_features["motion_variance"] * 10.0)
    
    # Duration-based total indicators
    total_indicators = int(max(2000, min(20000, duration * 300)))  # ~300 per second
    
    # Generate scores based on actual video features + some variation
    # Engaging: Higher motion with variety suggests engaging presence
    engaging_base = 0.40 + motion_factor * 0.30 + variance_factor * 0.15
    engaging_pos = int(total_indicators * float(np.clip(engaging_base + rng.uniform(-0.10, 0.10), 0.2, 0.8)))
    
    # Information sharing: Moderate motion, stable patterns
    info_base = 0.40 + motion_factor * 0.20 + (1 - variance_factor) * 0.10
    info_pos = int(total_indicators * float(np.clip(info_base + rng.uniform(-0.10, 0.10), 0.2, 0.8)))
    
    # Convincing: Strong, confident motion
    convince_base = 0.50 + motion_factor * 0.25
    convince_pos = int(total_indicators * float(np.clip(convince_base + rng.uniform(-0.10, 0.10), 0.3, 0.9)))
    
    # Authority: Controlled, deliberate motion (moderate intensity, low variance)
    authority_base = 0.30 + motion_factor * 0.20 - variance_factor * 0.10
    authority_pos = int(total_indicators * float(np.clip(authority_base + rng.uniform(-0.10, 0.10), 0.2, 0.7)))

    # Score Calculation Formula:
    # score = (positive_indicators / total_indicators) * 10
    # This converts the ratio of positive indicators to a 1-10 scale
    # Example: If 60% of indicators are positive, score = 0.6 * 10 = 6
    def pos_to_score(positive_count):
        """
        Converts positive indicator count to score (1-10).
        
        Args:
            positive_count: Number of positive indicators detected
            
        Returns:
            Integer score from 1 to 10
            
        Formula: (positive_count / total_indicators) * 10
        """
        ratio = positive_count / max(1, total_indicators)  # Prevent division by zero
        score = round(ratio * 10)  # Scale to 1-10
        return int(np.clip(score, 1, 10))  # Ensure score is between 1-10

    def score_to_scale(score):
        """
        Converts numeric score to descriptive scale.
        
        Args:
            score: Integer score from 1-10
            
        Returns:
            "low" (1-2), "moderate" (3-4), or "high" (5-10)
        """
        if score <= 2:
            return "low"
        if score <= 4:
            return "moderate"
        return "high"

    # Calculate scores for each category using the scoring formula
    engaging_score = pos_to_score(engaging_pos)
    info_score = pos_to_score(info_pos)
    convince_score = pos_to_score(convince_pos)
    authority_score = pos_to_score(authority_pos)

    # Overall Score Calculation:
    # Average of all four category scores, rounded to nearest integer
    # This gives equal weight to all categories
    overall_score = int(round(np.mean([engaging_score, info_score, convince_score, authority_score])))

    # Presence insights based on actual video features
    # Map motion features to presence characteristics
    motion_intensity = video_features["movement_intensity"]
    motion_stability = 1.0 - min(1.0, video_features["motion_variance"] * 5.0)
    
    # Eye contact: Higher motion suggests more engagement (but can vary)
    if motion_intensity > 0.3:
        eye_contact = rng.choice(["Moderate", "High"], p=[0.6, 0.4])
    elif motion_intensity > 0.15:
        eye_contact = rng.choice(["Low", "Moderate"], p=[0.3, 0.7])
    else:
        eye_contact = rng.choice(["Low", "Moderate"], p=[0.7, 0.3])
    
    # Stance: Lower variance suggests more stable stance
    if motion_stability > 0.6:
        stance = rng.choice(["Moderate", "Moderate–High", "High"], p=[0.3, 0.4, 0.3])
    elif motion_stability > 0.3:
        stance = rng.choice(["Low", "Moderate", "Moderate–High"], p=[0.2, 0.5, 0.3])
    else:
        stance = rng.choice(["Low", "Moderate"], p=[0.6, 0.4])
    
    # Uprightness: Based on motion patterns
    if motion_stability > 0.5 and motion_intensity > 0.2:
        upright = rng.choice(["Good", "Excellent"], p=[0.6, 0.4])
    elif motion_stability > 0.3:
        upright = rng.choice(["Fair", "Good"], p=[0.4, 0.6])
    else:
        upright = rng.choice(["Needs improvement", "Fair"], p=[0.5, 0.5])
    
    # AI-computed insights based on video features
    motion_variance = video_features["motion_variance"]
    if motion_variance < 0.02:
        gaze_stability = "Very stable"
    elif motion_variance < 0.05:
        gaze_stability = "Stable with occasional shifts"
    elif motion_variance < 0.1:
        gaze_stability = "Moderately stable"
    else:
        gaze_stability = "Unstable (frequent shifts)"
    
    if motion_intensity < 0.1:
        body_sway = "Very stable posture"
    elif motion_intensity < 0.2:
        body_sway = "Minimal sway"
    elif motion_intensity < 0.3:
        body_sway = "Moderate sway"
    else:
        body_sway = "Excessive sway detected"
    
    # Head tilt and shoulder symmetry based on motion patterns
    if motion_stability > 0.6:
        head_tilt = "Well-aligned"
    elif motion_stability > 0.4:
        head_tilt = "Generally upright"
    elif motion_stability > 0.2:
        head_tilt = "Slight tilt"
    else:
        head_tilt = "Noticeable tilt detected"
    
    if motion_stability > 0.7:
        shoulder_symmetry = "Excellent symmetry"
    elif motion_stability > 0.5:
        shoulder_symmetry = "Good symmetry"
    elif motion_stability > 0.3:
        shoulder_symmetry = "Moderately symmetric"
    else:
        shoulder_symmetry = "Asymmetric (uneven shoulders)"

    presence_comment = (
        f"Eye contact: {eye_contact}\n"
        f"Stance & grounding: {stance}\n"
        f"Uprightness & alignment: {upright}\n"
        f"Additional note: Gaze stability: {gaze_stability}. Body sway: {body_sway}. Head tilt: {head_tilt}. Shoulder symmetry: {shoulder_symmetry}."
    )

    # Load reference data for detection
    effort_df = load_effort_reference()
    shape_df = load_shape_reference()
    
    # Generate detection counts based on actual video features
    # Different videos will have different motion patterns
    effort_motions = effort_df['Motion Type'].tolist()
    shape_motions = shape_df['Motion Type'].tolist()
    
    # Create detection results dictionaries for graphs
    effort_detection = {}
    shape_detection = {}
    
    # Use video features to determine effort motion distribution
    total_frames = int(duration * 30)  # Assume 30 fps for detection sampling
    
    # Effort motions: Different motion intensities suggest different effort types
    # Based on Excel definitions:
    # Gliding: Smooth, continuous motion (moderate intensity, low variance) - Sustained, linear
    gliding_base = 0.12 + motion_factor * 0.18 - variance_factor * 0.10
    gliding_count = int(total_frames * np.clip(gliding_base + rng.uniform(-0.04, 0.04), 0.05, 0.35))
    
    # Punching: Strong, direct motion (high intensity) - Sudden, straight
    punching_base = 0.10 + motion_factor * 0.22
    punching_count = int(total_frames * np.clip(punching_base + rng.uniform(-0.04, 0.04), 0.05, 0.32))
    
    # Dabbing: Quick, precise motion (moderate intensity, medium variance) - Sudden, small
    dabbing_base = 0.10 + motion_factor * 0.12 + variance_factor * 0.05
    dabbing_count = int(total_frames * np.clip(dabbing_base + rng.uniform(-0.04, 0.04), 0.05, 0.30))
    
    # Flicking: Outward with quick rebound (moderate-high intensity, high variance) - Sudden, arced
    flicking_base = 0.08 + motion_factor * 0.15 + variance_factor * 0.12
    flicking_count = int(total_frames * np.clip(flicking_base + rng.uniform(-0.04, 0.04), 0.04, 0.28))
    
    # Pressing: Sustained, controlled motion (moderate intensity, low variance) - Sustained, linear
    pressing_base = 0.08 + motion_factor * 0.14 - variance_factor * 0.08
    pressing_count = int(total_frames * np.clip(pressing_base + rng.uniform(-0.04, 0.04), 0.04, 0.28))
    
    # Store effort detections
    effort_detection["Gliding"] = gliding_count
    effort_detection["gliding_count"] = gliding_count
    effort_detection["Punching"] = punching_count
    effort_detection["punching_count"] = punching_count
    effort_detection["Dabbing"] = dabbing_count
    effort_detection["dabbing_count"] = dabbing_count
    effort_detection["Flicking"] = flicking_count
    effort_detection["flicking_count"] = flicking_count
    effort_detection["Pressing"] = pressing_count
    effort_detection["pressing_count"] = pressing_count
    
    # Shape motions: Based on movement direction and pattern
    # Advancing: Forward motion (positive motion intensity)
    advancing_base = 0.10 + max(0, motion_factor - 0.3) * 0.20
    advancing_count = int(total_frames * np.clip(advancing_base + rng.uniform(-0.05, 0.05), 0.05, 0.30))
    
    # Retreating: Backward motion (can vary based on patterns)
    retreating_base = 0.08 + variance_factor * 0.15
    retreating_count = int(total_frames * np.clip(retreating_base + rng.uniform(-0.05, 0.05), 0.05, 0.30))
    
    # Enclosing: Inward motion (lower intensity, controlled)
    enclosing_base = 0.06 + (1 - motion_factor) * 0.10
    enclosing_count = int(total_frames * np.clip(enclosing_base + rng.uniform(-0.05, 0.05), 0.05, 0.25))
    
    # Spreading: Outward motion (moderate to high intensity)
    spreading_base = 0.12 + motion_factor * 0.18
    spreading_count = int(total_frames * np.clip(spreading_base + rng.uniform(-0.05, 0.05), 0.05, 0.35))
    
    # Directing: Focused, linear motion (moderate variance)
    directing_base = 0.10 + motion_factor * 0.12 + (1 - variance_factor) * 0.08
    directing_count = int(total_frames * np.clip(directing_base + rng.uniform(-0.05, 0.05), 0.05, 0.30))
    
    # Indirecting: Variable, shifting motion (high variance)
    indirecting_base = 0.08 + variance_factor * 0.20
    indirecting_count = int(total_frames * np.clip(indirecting_base + rng.uniform(-0.05, 0.05), 0.05, 0.35))
    
    # Store shape detections
    shape_detection["Advancing"] = advancing_count
    shape_detection["advancing_count"] = advancing_count
    shape_detection["Retreating"] = retreating_count
    shape_detection["retreating_count"] = retreating_count
    shape_detection["Enclosing"] = enclosing_count
    shape_detection["enclosing_count"] = enclosing_count
    shape_detection["Spreading"] = spreading_count
    shape_detection["spreading_count"] = spreading_count
    shape_detection["Directing"] = directing_count
    shape_detection["directing_count"] = directing_count
    shape_detection["Indirecting"] = indirecting_count
    shape_detection["indirecting_count"] = indirecting_count

    return {
        "duration_seconds": duration,
        "total_indicators": total_indicators,
        "engaging_pos": engaging_pos,
        "info_pos": info_pos,
        "convince_pos": convince_pos,
        "authority_pos": authority_pos,
        "engaging_score": engaging_score,
        "info_score": info_score,
        "convince_score": convince_score,
        "authority_score": authority_score,
        "overall_score": overall_score,
        "presence_comment": presence_comment,
        "effort_detection": effort_detection,
        "shape_detection": shape_detection,
    }


# =========================
# DOCX report generator
# =========================
def build_docx_report(report: ReportData, out_path: str, graph1_path: str = None, graph2_path: str = None):
    doc = Document()

    # Base font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    # Balanced spacing (readable, but still stable for 4-page layout)
    try:
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(2)
        style.paragraph_format.line_spacing = 1.10
    except Exception:
        pass
    try:
        bullet_style = doc.styles["List Bullet"]
        bullet_style.paragraph_format.space_before = Pt(0)
        bullet_style.paragraph_format.space_after = Pt(1)
        bullet_style.paragraph_format.line_spacing = 1.10
    except Exception:
        pass

    # Add header image to every page
    # (Linux deployments are case-sensitive; support both Header.png and header.png)
    _base_dir = os.path.dirname(__file__)
    header_image_candidates = [
        os.path.join(_base_dir, "Header.png"),
        os.path.join(_base_dir, "header.png"),
    ]
    header_image_path = next((p for p in header_image_candidates if os.path.exists(p)), None)
    if header_image_path:
        # Access the first section's header
        section = doc.sections[0]
        header = section.header
        header_paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Clear any existing content
        header_paragraph.clear()
        run = header_paragraph.add_run()
        # Add image with max width of 6.5 inches (typical page width minus margins)
        run.add_picture(header_image_path, width=Inches(6.5))

    # Title
    title = doc.add_paragraph("Presentation Analysis Report")
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(18)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph("")  # spacer

    # Client + date (match sample PDF label formatting)
    p = doc.add_paragraph()
    p.add_run("Client Name:     ").bold = True
    p.add_run(report.client_name if report.client_name else "—")

    p = doc.add_paragraph()
    p.add_run("Analysis Date:   ").bold = True
    p.add_run(report.analysis_date if report.analysis_date else "—")

    # Video information (match sample PDF)
    doc.add_paragraph("")
    h = doc.add_paragraph("Video Information")
    h.runs[0].bold = True
    h.runs[0].font.size = Pt(12)
    p = doc.add_paragraph()
    p.add_run("Duration: ").bold = True
    # The sample shows seconds; keep mm:ss too for clarity
    duration_seconds = None
    try:
        if report.video_length_str and ":" in report.video_length_str:
            mm_s, ss_s = report.video_length_str.strip().split(":", 1)
            duration_seconds = int(mm_s) * 60 + int(ss_s)
    except Exception:
        duration_seconds = None
    if duration_seconds is not None:
        p.add_run(f"{duration_seconds} seconds")
        p.add_run(f" ({report.video_length_str})")
    else:
        p.add_run(report.video_length_str if report.video_length_str else "—")

    # --- Page 2: Detailed analysis ---
    if report.categories:
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

        h = doc.add_paragraph("Detailed Presentation Analysis")
        h.runs[0].bold = True
        h.runs[0].font.size = Pt(12)
        try:
            h.paragraph_format.space_before = Pt(0)   # keep content closer to top
            h.paragraph_format.space_after = Pt(8)
            h.paragraph_format.line_spacing = 1.10
        except Exception:
            pass

        # Categories
        for cat in report.categories:
            # “o Category:” line like the sample
            cat_header = doc.add_paragraph(f"o {cat.name_en}:")
            cat_header.runs[0].bold = True
            try:
                cat_header.paragraph_format.space_before = Pt(8)
                cat_header.paragraph_format.space_after = Pt(2)
                cat_header.paragraph_format.line_spacing = 1.10
            except Exception:
                pass

            # Bullet traits from templates (English-only)
            tpl = REPORT_CATEGORY_TEMPLATES.get(cat.name_en)
            if tpl:
                bullets = tpl.get("bullets") or []
                if isinstance(bullets, list):
                    for b in bullets:
                        b = str(b).strip()
                        if b:
                            bp = doc.add_paragraph(b, style="List Bullet")
                            try:
                                bp.paragraph_format.space_before = Pt(0)
                                bp.paragraph_format.space_after = Pt(1)
                            except Exception:
                                pass

            # Scale + Description lines like the sample
            p = doc.add_paragraph()
            p.add_run("Scale: ").bold = True
            p.add_run(_scale_label(cat.scale))
            try:
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.line_spacing = 1.10
            except Exception:
                pass

            if cat.total > 0:
                p = doc.add_paragraph()
                p.add_run("Description: ").bold = True
                p.add_run(f"Detected {cat.positives} positive indicators out of {cat.total} total indicators")
                try:
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(6)
                    p.paragraph_format.line_spacing = 1.10
                except Exception:
                    pass

            # Avoid extra blank paragraphs that can cause paging differences
            # doc.add_paragraph("")

    # --- Page 3: Graph 1 (Effort) ---
    has_graph1 = bool(graph1_path and os.path.exists(graph1_path))
    has_graph2 = bool(graph2_path and os.path.exists(graph2_path))

    if has_graph1:
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

        h = doc.add_paragraph("Effort Motion Detection Results")
        h.runs[0].bold = True
        h.runs[0].font.size = Pt(12)

        graph1_paragraph = doc.add_paragraph()
        graph1_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = graph1_paragraph.add_run()
        # Slightly smaller to avoid Word pushing the content and creating a blank page
        run.add_picture(graph1_path, width=Inches(6.0))

    # --- Page 4: Graph 2 (Shape) ---
    if has_graph2:
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

        h = doc.add_paragraph("Shape Motion Detection Results")
        h.runs[0].bold = True
        h.runs[0].font.size = Pt(12)

        graph2_paragraph = doc.add_paragraph()
        graph2_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = graph2_paragraph.add_run()
        # Slightly smaller to avoid Word pushing the content and creating a blank page
        run.add_picture(graph2_path, width=Inches(6.0))

    # Add footer image and text to every page
    footer_image_path = os.path.join(os.path.dirname(__file__), "Footer.png")
    section = doc.sections[0]
    footer = section.footer
    
    # Clear existing footer paragraphs
    for paragraph in footer.paragraphs:
        paragraph.clear()
    if not footer.paragraphs:
        footer.add_paragraph()
    
    # Add footer image if it exists
    if os.path.exists(footer_image_path):
        footer_image_paragraph = footer.paragraphs[0]
        footer_image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_image_paragraph.add_run()
        # Add image with max width of 6.5 inches (typical page width minus margins)
        run.add_picture(footer_image_path, width=Inches(6.5))
        # Add a new paragraph for footer text
        footer_text_paragraph = footer.add_paragraph()
    else:
        footer_text_paragraph = footer.paragraphs[0]
    
    # Add footer text (generated by)
    footer_text_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer_text_paragraph.add_run(report.generated_by)
    footer_run.italic = True

    doc.save(out_path)


def build_detection_timeseries_excel_bytes(result: dict) -> bytes:
    """
    Build an Excel file (bytes) with one row per second and one column per motion.
    Excludes: Floating, Slashing, Wringing (not produced by this app's effort list anyway).
    """
    duration_seconds = float(result.get("duration_seconds") or 0.0)
    per_second = result.get("per_second_detection") or {}

    # Columns: all efforts + all shapes (explicit order)
    effort_cols = ["Gliding", "Punching", "Dabbing", "Flicking", "Pressing"]
    shape_cols = ["Advancing", "Retreating", "Enclosing", "Spreading", "Directing", "Indirecting"]
    cols = effort_cols + shape_cols

    n_secs = int(math.ceil(duration_seconds)) if duration_seconds > 0 else (max(per_second.keys()) + 1 if per_second else 0)
    rows = []
    for sec in range(n_secs):
        sec_data = per_second.get(sec, {})
        row = {"second": sec}
        for c in cols:
            row[c] = int(sec_data.get(c, 0))
        rows.append(row)

    df = pd.DataFrame(rows, columns=["second"] + cols)

    bio = io.BytesIO()
    with pd.ExcelWriter(bio, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="detections_per_second")
    return bio.getvalue()


# =========================
# Streamlit UI
# =========================
st.set_page_config(page_title="AI People Reader/Presentation Analysis", layout="wide")

# Custom CSS for elegant dark theme matching reference design
st.markdown("""
<style>
    /* Main container styling */
    .stApp {
        background-color: #3D3D37;
    }
    
    /* Text input and text area styling - beige rounded boxes */
    .stTextInput > div > div > input,
    .stTextArea > div > div > textarea {
        background-color: #C9BCA9 !important;
        color: #000000 !important;
        border-radius: 15px !important;
        border: none !important;
        padding: 15px !important;
        font-family: serif !important;
    }

    /* Placeholders inside light inputs should also be dark */
    .stTextInput input::placeholder,
    .stTextArea textarea::placeholder {
        color: #111111 !important;
        opacity: 0.7 !important;
    }
    
    /* Button styling - beige rounded */
    .stButton > button {
        background-color: #C9BCA9 !important;
        color: #000000 !important;
        border-radius: 15px !important;
        border: none !important;
        padding: 12px 30px !important;
        font-weight: bold !important;
        font-family: serif !important;
    }
    
    .stButton > button:hover {
        background-color: #B8AA98 !important;
    }
    
    /* File uploader styling */
    .stFileUploader > div {
        background-color: #52524A !important;
        border-radius: 15px !important;
        padding: 20px !important;
    }
    
    /* Slider styling */
    .stSlider > div > div > div {
        background-color: #C9BCA9 !important;
    }
    
    /* Selectbox and other inputs */
    .stSelectbox > div > div > div,
    .stCheckbox > label {
        color: #C9BCA9 !important;
        font-family: serif !important;
    }

    /* === Force BLACK text inside the beige/brown rectangle controls === */
    /* BaseWeb components used by Streamlit (selectboxes, inputs, etc.) */
    div[data-baseweb="select"] > div {
        background-color: #C9BCA9 !important;
        color: #000000 !important;
        border-radius: 15px !important;
        border: none !important;
    }
    div[data-baseweb="select"] * {
        color: #000000 !important;
    }
    div[data-baseweb="input"] input,
    div[data-baseweb="textarea"] textarea {
        color: #000000 !important;
    }

    /* File uploader "Browse files" button */
    section[data-testid="stFileUploader"] button {
        background-color: #C9BCA9 !important;
        color: #000000 !important;
        border-radius: 15px !important;
        border: none !important;
        font-weight: bold !important;
        font-family: serif !important;
    }
    section[data-testid="stFileUploader"] button * {
        color: #000000 !important;
    }
    
    /* Headers and title colors */
    h1, h2, h3, h4, h5, h6 {
        color: #C9BCA9 !important;
        font-family: serif !important;
    }
    
    /* General text color */
    p, label, .stMarkdown {
        color: #C9BCA9 !important;
        font-family: serif !important;
    }
    
    /* Expander styling */
    .streamlit-expanderHeader {
        background-color: #52524A !important;
        color: #C9BCA9 !important;
        border-radius: 10px !important;
    }
    
    /* Download button */
    .stDownloadButton > button {
        background-color: #C9BCA9 !important;
        color: #000000 !important;
        border-radius: 15px !important;
        font-weight: bold !important;
    }

    /* Status boxes (st.success / st.warning / st.error / st.info) text should be black */
    div[data-testid="stAlert"] * {
        color: #000000 !important;
    }
    
    /* Remove Streamlit's default spacing between elements */
    .element-container {
        margin-bottom: 0 !important;
        padding-bottom: 0 !important;
    }
</style>
""", unsafe_allow_html=True)

def _file_to_data_uri(path: str, mime: str = "image/png") -> str:
    with open(path, "rb") as f:
        b = f.read()
    return f"data:{mime};base64,{base64.b64encode(b).decode('ascii')}"


def _logo_png_to_data_uri(path: str) -> str:
    """
    Crop near-white padding around the logo, then encode to a PNG data-uri.
    This avoids the 'tiny logo' problem caused by huge whitespace in the source file.
    """
    bgr = cv2.imread(path, cv2.IMREAD_COLOR)
    if bgr is None:
        return _file_to_data_uri(path, mime="image/png")

    h, w = bgr.shape[:2]
    mask = np.any(bgr < 250, axis=2)
    ys, xs = np.where(mask)
    if ys.size > 0 and xs.size > 0:
        y1, y2 = int(ys.min()), int(ys.max())
        x1, x2 = int(xs.min()), int(xs.max())
        pad = 10
        y1 = max(0, y1 - pad)
        y2 = min(h - 1, y2 + pad)
        x1 = max(0, x1 - pad)
        x2 = min(w - 1, x2 + pad)
        bgr = bgr[y1 : y2 + 1, x1 : x2 + 1]

    ok, buf = cv2.imencode(".png", bgr)
    if not ok:
        return _file_to_data_uri(path, mime="image/png")
    return f"data:image/png;base64,{base64.b64encode(buf.tobytes()).decode('ascii')}"


# Render ONE combined hero image at the top (as requested)
_base_dir = os.path.dirname(__file__)
_combined_hero_path = os.path.join(_base_dir, "logo and header.png")
if os.path.exists(_combined_hero_path):
    hero_uri = _file_to_data_uri(_combined_hero_path, mime="image/png")
    st.markdown(
        f"""
<style>
  /* Make the top hero image full-bleed (edge-to-edge) like the reference site */
  section.main > div.block-container {{
    max-width: 100% !important;
    padding-left: 0 !important;
    padding-right: 0 !important;
  }}
  .stApp {{
    overflow-x: hidden;
  }}
  .pr-hero-one {{
    background: #3D3D37;
    padding: 0;
    display: flex;
    justify-content: center;
    width: 100vw;
    margin-left: calc(50% - 50vw);
  }}
  .pr-hero-one img {{
    width: 100%;
    max-width: none;
    height: auto;
    display: block;
  }}
</style>
<div class="pr-hero-one">
  <img src="{hero_uri}" alt="People Reader hero" />
</div>
""",
        unsafe_allow_html=True,
    )
else:
    st.warning('Missing required header image: "logo and header.png"')

st.markdown("<h1 style='white-space: nowrap;'>AI People Reader/Presentation Analysis</h1>", unsafe_allow_html=True)

st.write("Upload a video, then generate a DOCX report in your sample format.")

client_name = st.text_input("Client Name", value="")

uploaded_video = st.file_uploader("Upload presentation video", type=["mp4", "mov", "mkv", "avi"])

analysis_mode = st.selectbox(
    "Analysis Mode",
    options=[
        "Real (MediaPipe Pose + FaceMesh)",
        "Fallback (simple motion features + seeded randomness)",
    ],
    index=0 if mp is not None else 1,
)

# --- Auto recommend settings (useful when analyzing many videos) ---
auto_recommend = st.checkbox("Auto recommend fps/max-frames for this uploaded video", value=True)
auto_apply = st.checkbox("Auto-apply recommended values on upload", value=True, disabled=not auto_recommend)
auto_apply_advanced = st.checkbox("Auto-apply advanced accuracy settings", value=True, disabled=not auto_recommend)
recommend_preset = st.selectbox(
    "Recommendation preset",
    options=["Fast", "Balanced", "Accurate"],
    index=1,
    disabled=not auto_recommend,
    help="Fast = quickest (lower fps/cap), Balanced = default, Accurate = slowest but most coverage.",
)

# Compute recommended settings (needs a temp file because OpenCV reads from disk)
rec_sample_fps = None
rec_max_frames = None
rec_duration = None
if uploaded_video is not None and auto_recommend:
    try:
        # Detect upload changes without reading the entire file into memory.
        # (Large uploads can otherwise cause high memory usage and backend restarts on small servers.)
        upload_sig = f"{uploaded_video.name}:{getattr(uploaded_video, 'size', 'unknown')}"
        preset_sig = (recommend_preset or "Balanced").strip()
        rec_sig = f"{upload_sig}:{preset_sig}"

        if st.session_state.get("_upload_sig") != rec_sig:
            # New upload: write temp file + compute duration
            st.session_state["_upload_sig"] = rec_sig
            with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(uploaded_video.name)[1]) as tmp:
                try:
                    uploaded_video.seek(0)
                except Exception:
                    pass
                shutil.copyfileobj(uploaded_video, tmp)
                st.session_state["_upload_tmp_path"] = tmp.name

            rec_duration = get_video_duration_seconds(st.session_state["_upload_tmp_path"])
            st.session_state["_upload_duration_seconds"] = float(rec_duration)

            r_fps, r_frames = recommend_sampling_settings(float(rec_duration), preset=recommend_preset)
            st.session_state["_rec_sample_fps"] = int(r_fps)
            st.session_state["_rec_max_frames"] = int(r_frames)

            if auto_apply:
                st.session_state["sample_fps"] = int(r_fps)
                st.session_state["max_frames"] = int(r_frames)

            if auto_apply_advanced:
                adv = recommend_advanced_settings(recommend_preset)
                st.session_state["pose_model_complexity"] = int(adv["pose_model_complexity"])
                st.session_state["pose_min_detection_confidence"] = float(adv["pose_min_detection_confidence"])
                st.session_state["pose_min_tracking_confidence"] = float(adv["pose_min_tracking_confidence"])
                st.session_state["face_min_detection_confidence"] = float(adv["face_min_detection_confidence"])
                st.session_state["facemesh_min_detection_confidence"] = float(adv["facemesh_min_detection_confidence"])
                st.session_state["facemesh_min_tracking_confidence"] = float(adv["facemesh_min_tracking_confidence"])
        else:
            rec_duration = float(st.session_state.get("_upload_duration_seconds") or 0.0)

        rec_sample_fps = int(st.session_state.get("_rec_sample_fps") or 5)
        rec_max_frames = int(st.session_state.get("_rec_max_frames") or 300)

        coverage_seconds = (rec_max_frames / rec_sample_fps) if rec_sample_fps else 0.0
        st.caption(
            f"Recommended ({recommend_preset}): **{rec_sample_fps} fps**, **{rec_max_frames} frames** "
            f"(≈ {coverage_seconds:.0f}s coverage at that sampling rate; video duration ≈ {rec_duration:.0f}s)"
        )

        if not auto_apply:
            if st.button("Apply recommended settings"):
                st.session_state["sample_fps"] = int(rec_sample_fps)
                st.session_state["max_frames"] = int(rec_max_frames)
                st.rerun()
    except Exception:
        # If anything goes wrong, fall back to manual settings
        pass

sample_fps = st.slider("Analysis sampling rate (frames/sec)", min_value=1, max_value=10, value=5, key="sample_fps")
max_frames = st.slider("Max frames to analyze", min_value=50, max_value=1500, value=300, step=50, key="max_frames")

with st.expander("Advanced accuracy settings (slower but more accurate)"):
    pose_model_complexity = st.selectbox("Pose model complexity", options=[0, 1, 2], index=1, key="pose_model_complexity")
    pose_min_detection_confidence = st.slider("Pose min detection confidence", 0.1, 0.99, 0.5, 0.01, key="pose_min_detection_confidence")
    pose_min_tracking_confidence = st.slider("Pose min tracking confidence", 0.1, 0.99, 0.5, 0.01, key="pose_min_tracking_confidence")
    face_min_detection_confidence = st.slider("Face detection min confidence", 0.1, 0.99, 0.5, 0.01, key="face_min_detection_confidence")
    facemesh_min_detection_confidence = st.slider("FaceMesh min detection confidence", 0.1, 0.99, 0.5, 0.01, key="facemesh_min_detection_confidence")
    facemesh_min_tracking_confidence = st.slider("FaceMesh min tracking confidence", 0.1, 0.99, 0.5, 0.01, key="facemesh_min_tracking_confidence")

if mp is None and analysis_mode.startswith("Real"):
    st.warning("MediaPipe is not installed in this environment. Please install `mediapipe` to enable real analysis.")

st.divider()

custom_summary_prefix = st.text_area(
    "Optional: Summary/comment to include in the report",
    value="",
    height=120,
    placeholder="Type any summary/comment you want shown in the report (optional). Leave empty to omit the Summary section.",
)

st.subheader("Dot Motion (same upload)")
st.write("Optionally generate a dot-motion visualization video from the same uploaded file.")
make_dot_video = st.checkbox("Also generate dot-motion video", value=True)
dot_size = st.slider(
    "Dot Size (radius in pixels)",
    min_value=1,
    max_value=5,
    value=2,
    help="Smaller dots (1-2) match the reference image better",
    key="dot_size_slider",
)

generate = st.button("Generate Report")

if generate:
    if uploaded_video is None:
        st.error("Please upload a video first.")
        st.stop()

    with st.spinner("Analyzing video and generating report..."):
        # Save uploaded file to temp
        with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(uploaded_video.name)[1]) as tmp:
            try:
                uploaded_video.seek(0)
            except Exception:
                pass
            shutil.copyfileobj(uploaded_video, tmp)
            video_path = tmp.name

        # Analyze
        try:
            if analysis_mode.startswith("Real"):
                result = analyze_video_mediapipe(
                    video_path=video_path,
                    sample_fps=float(sample_fps),
                    max_frames=int(max_frames),
                    pose_model_complexity=int(pose_model_complexity),
                    pose_min_detection_confidence=float(pose_min_detection_confidence),
                    pose_min_tracking_confidence=float(pose_min_tracking_confidence),
                    face_min_detection_confidence=float(face_min_detection_confidence),
                    facemesh_min_detection_confidence=float(facemesh_min_detection_confidence),
                    facemesh_min_tracking_confidence=float(facemesh_min_tracking_confidence),
                )
            else:
                result = analyze_video_placeholder(video_path=video_path, seed=42)
        except Exception as e:
            st.warning(f"Real analysis failed ({str(e)}). Falling back to placeholder analysis.")
            result = analyze_video_placeholder(video_path=video_path, seed=42)

        duration_str = format_seconds_to_mmss(result["duration_seconds"])
        analysis_date = datetime.now().strftime("%Y-%m-%d")

        total = int(result["total_indicators"])

        categories = [
            CategoryResult(
                name_en="Engaging & Connecting",
                name_th="",
                score=int(result["engaging_score"]),
                scale=("moderate" if int(result["engaging_score"]) in [3, 4] else ("high" if int(result["engaging_score"]) >= 5 else "low")),
                positives=int(result["engaging_pos"]),
                total=total,
            ),
            CategoryResult(
                name_en="Confidence",
                name_th="",
                score=int(result["convince_score"]),
                scale=("moderate" if int(result["convince_score"]) in [3, 4] else ("high" if int(result["convince_score"]) >= 5 else "low")),
                positives=int(result["convince_pos"]),
                total=total,
            ),
            CategoryResult(
                name_en="Authority",
                name_th="",
                score=int(result["authority_score"]),
                scale=("moderate" if int(result["authority_score"]) in [3, 4] else ("high" if int(result["authority_score"]) >= 5 else "low")),
                positives=int(result["authority_pos"]),
                total=total,
            ),
        ]

        # Summary section: user-provided only (do not include AI presence notes)
        summary_comment = custom_summary_prefix.strip()

        report = ReportData(
            client_name=client_name.strip(),
            analysis_date=analysis_date,
            video_length_str=duration_str,
            # Overall score should reflect what is actually shown in the report
            overall_score=int(round(float(np.mean([c.score for c in categories])))) if categories else 0,
            categories=categories,
            summary_comment=summary_comment,
            generated_by="Generated by AI People Reader™",
        )

        # Build docx
        out_dir = tempfile.mkdtemp()
        out_path = os.path.join(out_dir, f"Presentation_Analysis_Report_{analysis_date}.docx")
        
        # Generate graphs based on detection results
        graph1_path = os.path.join(out_dir, "Graph 1.png")
        graph2_path = os.path.join(out_dir, "Graph 2.png")
        
        # Also save graphs to app directory
        app_dir = os.path.dirname(__file__)
        graph1_app_path = os.path.join(app_dir, "Graph 1.png")
        graph2_app_path = os.path.join(app_dir, "Graph 2.png")
        
        try:
            generate_effort_graph(result.get("effort_detection", {}), result.get("shape_detection", {}), graph1_path)
            generate_shape_graph(result.get("shape_detection", {}), graph2_path)
            
            # Copy to app directory
            shutil.copy2(graph1_path, graph1_app_path)
            shutil.copy2(graph2_path, graph2_app_path)
        except Exception as e:
            st.warning(f"Graph generation had issues: {str(e)}")
            # Create empty graphs if generation fails
            if not os.path.exists(graph1_path):
                plt.figure(figsize=(10, 6))
                plt.text(0.5, 0.5, "Graph generation failed", ha='center', va='center')
                plt.savefig(graph1_path)
                plt.close()
            if not os.path.exists(graph2_path):
                plt.figure(figsize=(10, 6))
                plt.text(0.5, 0.5, "Graph generation failed", ha='center', va='center')
                plt.savefig(graph2_path)
                plt.close()
        
        build_docx_report(report, out_path, graph1_path=graph1_path, graph2_path=graph2_path)

        # Offer download
        with open(out_path, "rb") as f:
            docx_bytes = f.read()

        st.success("Report generated!")
        st.download_button(
            label="Download DOCX Report",
            data=docx_bytes,
            file_name=os.path.basename(out_path),
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        # Dot video generation (same upload)
        if make_dot_video:
            if mp is None:
                st.warning("MediaPipe is not installed, so dot motion generation is unavailable. Install `mediapipe` to enable it.")
            else:
                dot_progress = st.progress(0.0)
                dot_status = st.empty()

                def _dot_progress(pct: float, cur: int, total: int):
                    dot_progress.progress(float(max(0.0, min(1.0, pct))))
                    dot_status.text(f"Dot video: processing frame {cur} of {total}")

                try:
                    dot_mp4_bytes, dot_preview_frame = generate_dot_motion_video_bytes(
                        input_video_path=video_path,
                        dot_size=int(dot_size),
                        progress_cb=_dot_progress,
                    )
                    dot_progress.empty()
                    dot_status.empty()

                    st.download_button(
                        label="Download Dot Motion Video (MP4)",
                        data=dot_mp4_bytes,
                        file_name="dot_motion_output.mp4",
                        mime="video/mp4",
                    )

                    if dot_preview_frame is not None:
                        st.image(dot_preview_frame, caption="Dot motion preview frame", use_container_width=True)

                    if VideoFileClip is None:
                        st.caption("Note: MoviePy is not installed, so the dot video may not include the original audio.")
                except Exception as e:
                    dot_progress.empty()
                    dot_status.empty()
                    st.warning(f"Dot motion generation failed: {str(e)}")

        # Offer detection Excel download (one row per second)
        try:
            excel_bytes = build_detection_timeseries_excel_bytes(result)
            st.download_button(
                label="Download Detection Excel (per second)",
                data=excel_bytes,
                file_name=f"Detection_Per_Second_{analysis_date}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )
        except Exception as e:
            st.warning(f"Could not build detection Excel: {str(e)}")

        # Quick preview of key outputs
        st.subheader("Preview (key fields)")
        st.write(
            {
                "Client Name": report.client_name,
                "Analysis Date": report.analysis_date,
                "Video Length": report.video_length_str,
                "Overall Score": report.overall_score,
            }
        )

        st.subheader("Category scores")
        st.dataframe(pd.DataFrame([{
            "Category": c.name_en if not c.name_th else f"{c.name_en} ({c.name_th})",
            "Score": c.score,
            "Scale": c.scale,
            "Positive": c.positives,
            "Total": c.total,
        } for c in categories]))

        # Cleanup temp video (must happen AFTER dot motion generation, which reads video_path)
        try:
            os.remove(video_path)
        except Exception:
            pass
