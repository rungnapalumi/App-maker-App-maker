# src/report_worker.py
# ------------------------------------------------------------
# AI People Reader - Report Worker (TH/EN)  [LEGACY QUEUE]
# ------------------------------------------------------------

import os
import io
import json
import time
import shutil
import tempfile
import logging
import subprocess
from dataclasses import dataclass
from datetime import datetime, timezone
from typing import Optional, Tuple, Dict, Any, List

import boto3

# -------------------------
# Optional heavy libs (safe import)
# -------------------------
try:
    import cv2  # type: ignore
except Exception:
    cv2 = None  # type: ignore

try:
    import numpy as np  # type: ignore
except Exception:
    np = None  # type: ignore

try:
    import pandas as pd  # type: ignore
except Exception:
    pd = None  # type: ignore

try:
    import mediapipe as mp  # type: ignore
    MP_HAS_SOLUTIONS = hasattr(mp, "solutions")
except Exception:
    mp = None  # type: ignore
    MP_HAS_SOLUTIONS = False

# Headless matplotlib (Render-safe)
try:
    import matplotlib  # type: ignore
    matplotlib.use("Agg")
except Exception:
    pass

try:
    import matplotlib.pyplot as plt  # type: ignore
except Exception:
    plt = None  # type: ignore

try:
    from docx import Document  # type: ignore
    from docx.shared import Pt, Inches  # type: ignore
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK  # type: ignore
except Exception:
    Document = None  # type: ignore
    Pt = Inches = None  # type: ignore
    WD_ALIGN_PARAGRAPH = WD_BREAK = None  # type: ignore


# -------------------------
# Logging
# -------------------------
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s | %(levelname)s | %(message)s"
)
log = logging.getLogger("report_worker")


# -------------------------
# Paths / Assets
# -------------------------
PROJECT_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))

ASSET_HEADER = os.path.join(PROJECT_ROOT, "Header.png")
ASSET_FOOTER = os.path.join(PROJECT_ROOT, "Footer.png")
ASSET_EFFORT = os.path.join(PROJECT_ROOT, "Effort.xlsx")
ASSET_SHAPE = os.path.join(PROJECT_ROOT, "Shape.xlsx")

DEFAULT_MOVENET_MODEL = os.path.join(
    PROJECT_ROOT,
    "movenet_singlepose_lightning.tflite"
)
MOVENET_MODEL_PATH = os.getenv("MOVENET_MODEL_PATH") or DEFAULT_MOVENET_MODEL

NORMALIZE_FPS = int(os.getenv("REPORT_NORMALIZE_FPS", "30"))
POSE_SAMPLE_FPS = float(os.getenv("REPORT_POSE_SAMPLE_FPS", "6.0"))


# -------------------------
# S3 config
# -------------------------
AWS_BUCKET = os.getenv("AWS_BUCKET") or os.getenv("S3_BUCKET")
AWS_REGION = os.getenv("AWS_REGION", "ap-southeast-1")

JOB_POLL_INTERVAL = int(os.getenv("JOB_POLL_INTERVAL", "10"))

JOBS_PENDING_PREFIX = os.getenv("JOBS_PENDING_PREFIX", "jobs/pending/")
JOBS_PROCESSING_PREFIX = os.getenv("JOBS_PROCESSING_PREFIX", "jobs/processing/")
JOBS_FINISHED_PREFIX = os.getenv("JOBS_FINISHED_PREFIX", "jobs/finished/")
JOBS_FAILED_PREFIX = os.getenv("JOBS_FAILED_PREFIX", "jobs/failed/")

if not AWS_BUCKET:
    raise RuntimeError("Missing AWS_BUCKET or S3_BUCKET env var")

s3 = boto3.client("s3", region_name=AWS_REGION)


# =========================
# Text helpers
# =========================
def _t(lang: str, en: str, th: str) -> str:
    lang = (lang or "en").lower()
    return th if lang.startswith("th") else en


def _scale_label(scale: str, lang: str) -> str:
    s = (scale or "").lower()
    if s.startswith("high"):
        return "สูง" if lang.startswith("th") else "High"
    if s.startswith("mod"):
        return "ปานกลาง" if lang.startswith("th") else "Moderate"
    if s.startswith("low"):
        return "ต่ำ" if lang.startswith("th") else "Low"
    return "—"


# =========================
# Category templates
# =========================
REPORT_CATEGORY_TEMPLATES = {
    "Engaging & Connecting": {
        "bullets_en": [
            "Approachability",
            "Relatability",
            "Engagement and instant rapport",
        ],
        "bullets_th": [
            "ความเป็นกันเอง",
            "ความเข้าถึงได้",
            "การมีส่วนร่วมและสร้างความคุ้นเคยได้รวดเร็ว",
        ],
    },
    "Confidence": {
        "bullets_en": [
            "Optimistic presence",
            "Focus",
            "Ability to persuade",
        ],
        "bullets_th": [
            "บุคลิกเชิงบวก",
            "ความมีสมาธิ",
            "ความสามารถในการโน้มน้าว",
        ],
    },
    "Authority": {
        "bullets_en": [
            "Sense of importance",
            "Pressing for action",
        ],
        "bullets_th": [
            "ความสำคัญของประเด็น",
            "การผลักดันให้เกิดการลงมือทำ",
        ],
    },
}


# =========================
# Data models
# =========================
@dataclass
class CategoryResult:
    name_en: str
    name_th: str
    score: int
    scale: str
    positives: int
    total: int
    description: str = ""


@dataclass
class FirstImpressionItem:
    title_en: str
    title_th: str
    scale: str
    score: int
    insight_en: str
    insight_th: str
    impact_en: str
    impact_th: str
    details: Dict[str, Any]


@dataclass
class FirstImpressionResult:
    eye_contact: FirstImpressionItem
    uprightness: FirstImpressionItem
    stance: FirstImpressionItem
    reliability_note_en: str = ""
    reliability_note_th: str = ""


@dataclass
class ReportData:
    client_name: str
    analysis_date: str
    video_length_str: str
    overall_score: int
    first_impression: Optional[FirstImpressionResult]
    categories: list
    summary_comment: str
    generated_by: str


# =========================
# Video helpers
# =========================
def format_seconds_to_mmss(sec: float) -> str:
    sec = max(0.0, float(sec))
    m = int(sec // 60)
    s = int(round(sec - m * 60))
    if s == 60:
        m += 1
        s = 0
    return f"{m:02d}:{s:02d}"


def get_video_duration_seconds(path: str) -> float:
    if cv2 is None:
        return 0.0
    cap = cv2.VideoCapture(path)
    if not cap.isOpened():
        return 0.0
    fps = cap.get(cv2.CAP_PROP_FPS) or 0.0
    frames = cap.get(cv2.CAP_PROP_FRAME_COUNT) or 0.0
    cap.release()
    if fps <= 0:
        return 0.0
    return frames / fps
# =========================
# FFmpeg normalize (always)
# =========================
def ensure_ffmpeg() -> str:
    ff = shutil.which("ffmpeg")
    if not ff:
        raise RuntimeError("FFmpeg not found in PATH")
    return ff


def normalize_video_with_ffmpeg(
    input_path: str,
    output_path: str,
    fps: int = 30,
) -> None:
    ff = ensure_ffmpeg()
    cmd = [
        ff, "-y",
        "-i", input_path,
        "-vf", f"fps={fps},format=yuv420p",
        "-c:v", "libx264",
        "-preset", "veryfast",
        "-crf", "23",
        "-movflags", "+faststart",
        "-an",
        output_path,
    ]
    p = subprocess.run(
        cmd,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
    )

    # Render-safe: allow soft fail if output exists
    if p.returncode != 0:
        if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
            log.warning("FFmpeg returned non-zero but output exists; continuing")
            return
        raise RuntimeError(f"FFmpeg failed:\n{p.stderr[-1500:]}")


# =========================
# Excel loaders (Effort / Shape)
# =========================
def load_effort_reference(path: str = ASSET_EFFORT):
    if pd is None:
        raise RuntimeError("pandas required for Effort.xlsx")
    if not os.path.exists(path):
        raise FileNotFoundError(f"Effort.xlsx not found: {path}")

    df = pd.read_excel(path, header=None)
    df.columns = [
        "Motion Type",
        "Direction",
        "Body Part Involvement",
        "Pathway",
        "Timing",
        "Other Motion Clues",
    ]
    return df.iloc[2:].reset_index(drop=True)


def load_shape_reference(path: str = ASSET_SHAPE):
    if pd is None:
        raise RuntimeError("pandas required for Shape.xlsx")
    if not os.path.exists(path):
        raise FileNotFoundError(f"Shape.xlsx not found: {path}")

    df = pd.read_excel(path, header=None)
    df.columns = [
        "Motion Type",
        "Direction",
        "Body Part Involvement",
        "Pathway",
        "Timing",
        "Other Motion Clues",
    ]
    return df.iloc[2:].reset_index(drop=True)


# =========================
# Graph generators
# =========================
def generate_effort_graph(
    effort_detection: dict,
    shape_detection: dict,
    output_path: str,
):
    if plt is None or np is None:
        raise RuntimeError("matplotlib/numpy required")

    effort_df = load_effort_reference()
    shape_df = load_shape_reference()

    effort_motions = [
        m for m in effort_df["Motion Type"].tolist()
        if m not in ("Floating", "Slashing", "Wringing")
    ]
    shape_motions = shape_df["Motion Type"].tolist()

    motions = effort_motions + shape_motions
    counts = []

    for m in motions:
        c = (
            effort_detection.get(m)
            or effort_detection.get(f"{m.lower()}_count")
            or shape_detection.get(m)
            or shape_detection.get(f"{m.lower()}_count")
            or 0
        )
        counts.append(int(c))

    total = sum(counts) or 1
    pcts = [(c / total) * 100 for c in counts]

    sorted_data = sorted(zip(motions, pcts), key=lambda x: x[1], reverse=True)
    motions, pcts = zip(*sorted_data)

    fig, ax = plt.subplots(figsize=(14, max(7, len(motions) * 0.45)))
    bars = ax.barh(motions, pcts)

    ax.set_xlim(0, 100)
    ax.set_xlabel("Percentage (%)", fontweight="bold")
    ax.set_title("Effort Summary", fontweight="bold")
    ax.grid(axis="x", alpha=0.3)
    ax.invert_yaxis()

    for bar in bars:
        w = bar.get_width()
        if w > 0:
            ax.text(
                w + 1,
                bar.get_y() + bar.get_height() / 2,
                f"{w:.1f}%",
                va="center",
                fontweight="bold",
            )

    plt.tight_layout()
    plt.savefig(output_path, dpi=300, bbox_inches="tight")
    plt.close()


def generate_shape_graph(
    detection_results: dict,
    output_path: str,
):
    if plt is None or np is None:
        raise RuntimeError("matplotlib/numpy required")

    shape_df = load_shape_reference()
    motions = shape_df["Motion Type"].tolist()

    counts = [
        int(
            detection_results.get(m)
            or detection_results.get(f"{m.lower()}_count")
            or 0
        )
        for m in motions
    ]

    total = sum(counts) or 1
    pcts = [(c / total) * 100 for c in counts]

    fig, ax = plt.subplots(figsize=(max(12, len(motions) * 1.5), 7))
    bars = ax.bar(motions, pcts)

    ax.set_ylim(0, min(100, max(pcts) * 1.15 if pcts else 100))
    ax.set_ylabel("Percentage (%)", fontweight="bold")
    ax.set_title("Shape Motion Detection Results", fontweight="bold")
    ax.grid(axis="y", alpha=0.3)

    for bar in bars:
        h = bar.get_height()
        if h > 0:
            ax.text(
                bar.get_x() + bar.get_width() / 2,
                h,
                f"{h:.1f}%",
                ha="center",
                va="bottom",
                fontweight="bold",
            )

    plt.xticks(rotation=15, ha="right")
    plt.tight_layout()
    plt.savefig(output_path, dpi=300, bbox_inches="tight")
    plt.close()
# =========================
# First Impression — REAL analysis from video
# Primary: MoveNet TFLite
# Fallback: MediaPipe Pose
# =========================

# MoveNet keypoints (17):
# 0 nose, 1 left_eye, 2 right_eye, 3 left_ear, 4 right_ear,
# 5 left_shoulder, 6 right_shoulder, 7 left_elbow, 8 right_elbow,
# 9 left_wrist, 10 right_wrist, 11 left_hip, 12 right_hip,
# 13 left_knee, 14 right_knee, 15 left_ankle, 16 right_ankle
POSE_IDX = {
    "NOSE": 0,
    "LEFT_EAR": 3,
    "RIGHT_EAR": 4,
    "LEFT_SHOULDER": 5,
    "RIGHT_SHOULDER": 6,
    "LEFT_HIP": 11,
    "RIGHT_HIP": 12,
    "LEFT_ANKLE": 15,
    "RIGHT_ANKLE": 16,
}

# MediaPipe Pose landmark indices (BlazePose)
MP_IDX = {
    "NOSE": 0,
    "LEFT_EAR": 7,
    "RIGHT_EAR": 8,
    "LEFT_SHOULDER": 11,
    "RIGHT_SHOULDER": 12,
    "LEFT_HIP": 23,
    "RIGHT_HIP": 24,
    "LEFT_ANKLE": 27,
    "RIGHT_ANKLE": 28,
}


def _load_tflite_interpreter(model_path: str):
    try:
        from tflite_runtime.interpreter import Interpreter  # type: ignore
    except Exception:
        try:
            from tensorflow.lite.python.interpreter import Interpreter  # type: ignore
        except Exception as e:
            raise RuntimeError(f"tflite interpreter not available: {e}")

    if not os.path.exists(model_path):
        raise FileNotFoundError(
            f"MoveNet model not found at: {model_path}. "
            f"Set env MOVENET_MODEL_PATH or include the .tflite file in the repo."
        )

    interpreter = Interpreter(model_path=model_path)
    interpreter.allocate_tensors()
    return interpreter


def _movenet_infer(interpreter, frame_bgr: "np.ndarray") -> "np.ndarray":
    """
    Returns keypoints in shape (17, 3) with columns [x, y, score] in normalized coords.
    MoveNet output is [y, x, score] per keypoint -> we convert to [x,y,score]
    """
    if np is None or cv2 is None:
        raise RuntimeError("numpy/opencv required for MoveNet inference")

    input_details = interpreter.get_input_details()
    output_details = interpreter.get_output_details()

    rgb = cv2.cvtColor(frame_bgr, cv2.COLOR_BGR2RGB)

    in_h = int(input_details[0]["shape"][1])
    in_w = int(input_details[0]["shape"][2])
    resized = cv2.resize(rgb, (in_w, in_h), interpolation=cv2.INTER_AREA)

    in_dtype = input_details[0]["dtype"]
    x = resized
    if str(in_dtype).endswith("uint8"):
        x = x.astype(np.uint8)
    elif str(in_dtype).endswith("int32"):
        x = x.astype(np.int32)
    else:
        x = x.astype(np.float32) / 255.0

    x = np.expand_dims(x, axis=0)

    interpreter.set_tensor(input_details[0]["index"], x)
    interpreter.invoke()

    out = interpreter.get_tensor(output_details[0]["index"])

    if out.ndim == 4:
        out = out[0, 0, :, :]
    elif out.ndim == 3:
        out = out[0, :, :]
    else:
        raise RuntimeError(f"Unexpected MoveNet output shape: {out.shape}")

    # out: (17,3) [y,x,score]
    y = out[:, 0]
    xx = out[:, 1]
    s = out[:, 2]
    kps = np.stack([xx, y, s], axis=1)  # -> [x,y,score]
    return kps


def _fi_clamp01(x: float) -> float:
    return max(0.0, min(1.0, float(x)))


def _fi_level_from_score(score: int, lang: str) -> str:
    lang = (lang or "en").lower()
    if score >= 75:
        return "สูง" if lang.startswith("th") else "High"
    if score >= 45:
        return "ปานกลาง" if lang.startswith("th") else "Moderate"
    return "ต่ำ" if lang.startswith("th") else "Low"


def _fi_safe_mean(vals: List[float]) -> Optional[float]:
    if np is None or not vals:
        return None
    return float(np.mean(vals))


def _fi_vis_ok(p: Optional[Tuple[float, float, float]], thr: float = 0.35) -> bool:
    return bool(p) and float(p[2]) >= thr


def _fi_get(pose: Dict[str, Tuple[float, float, float]], k: str) -> Optional[Tuple[float, float, float]]:
    return pose.get(k)


def _fi_extract_pose_sequence_movenet(
    video_path: str,
    model_path: str,
    max_frames: int = 1000,
    sample_fps: float = 6.0,
) -> Tuple[List[Dict[str, Tuple[float, float, float]]], Dict[str, Any]]:
    if cv2 is None or np is None:
        return [], {"reason": "missing_libs"}

    cap = cv2.VideoCapture(video_path)
    if not cap.isOpened():
        return [], {"reason": "cannot_open_video"}

    fps = cap.get(cv2.CAP_PROP_FPS) or 25.0
    frame_count = int(cap.get(cv2.CAP_PROP_FRAME_COUNT) or 0)
    duration_sec = frame_count / fps if fps > 0 else 0.0

    if fps <= 0:
        fps = 25.0
    sample_every = int(max(1, round(fps / max(1e-6, sample_fps))))

    interpreter = _load_tflite_interpreter(model_path)

    seq: List[Dict[str, Tuple[float, float, float]]] = []
    idx = 0
    processed = 0

    while True:
        ok, frame = cap.read()
        if not ok:
            break
        idx += 1
        if idx % sample_every != 0:
            continue

        kps = _movenet_infer(interpreter, frame)  # (17,3) [x,y,score]

        pose_dict: Dict[str, Tuple[float, float, float]] = {}
        for name, i in POSE_IDX.items():
            if 0 <= i < kps.shape[0]:
                pose_dict[name] = (float(kps[i, 0]), float(kps[i, 1]), float(kps[i, 2]))
        seq.append(pose_dict)

        processed += 1
        if processed >= max_frames:
            break

    cap.release()

    meta = {
        "engine": "movenet_tflite",
        "fps": float(fps),
        "frame_count": int(frame_count),
        "duration_sec": float(duration_sec),
        "sample_every": int(sample_every),
        "processed_frames": int(len(seq)),
        "model_path": model_path,
    }
    return seq, meta


def _fi_extract_pose_sequence_mediapipe(
    video_path: str,
    max_frames: int = 1000,
    sample_fps: float = 6.0,
) -> Tuple[List[Dict[str, Tuple[float, float, float]]], Dict[str, Any]]:
    if cv2 is None or np is None or not MP_HAS_SOLUTIONS:
        return [], {"reason": "missing_libs_or_mediapipe"}

    cap = cv2.VideoCapture(video_path)
    if not cap.isOpened():
        return [], {"reason": "cannot_open_video"}

    fps = cap.get(cv2.CAP_PROP_FPS) or 25.0
    frame_count = int(cap.get(cv2.CAP_PROP_FRAME_COUNT) or 0)
    duration_sec = frame_count / fps if fps > 0 else 0.0

    if fps <= 0:
        fps = 25.0
    sample_every = int(max(1, round(fps / max(1e-6, sample_fps))))

    pose = mp.solutions.pose.Pose(
        static_image_mode=False,
        model_complexity=1,
        enable_segmentation=False,
        smooth_landmarks=True,
        min_detection_confidence=0.5,
        min_tracking_confidence=0.5,
    )

    seq: List[Dict[str, Tuple[float, float, float]]] = []
    idx = 0
    processed = 0

    while True:
        ok, frame = cap.read()
        if not ok:
            break
        idx += 1
        if idx % sample_every != 0:
            continue

        rgb = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
        res = pose.process(rgb)

        pose_dict: Dict[str, Tuple[float, float, float]] = {}
        if res and res.pose_landmarks and res.pose_landmarks.landmark:
            lm = res.pose_landmarks.landmark
            for name, i in MP_IDX.items():
                if 0 <= i < len(lm):
                    pose_dict[name] = (float(lm[i].x), float(lm[i].y), float(lm[i].visibility))

        seq.append(pose_dict)

        processed += 1
        if processed >= max_frames:
            break

    cap.release()
    try:
        pose.close()
    except Exception:
        pass

    meta = {
        "engine": "mediapipe_pose",
        "fps": float(fps),
        "frame_count": int(frame_count),
        "duration_sec": float(duration_sec),
        "sample_every": int(sample_every),
        "processed_frames": int(len(seq)),
    }
    return seq, meta


def _analyze_eye_contact(seq: List[Dict[str, Tuple[float, float, float]]], lang: str) -> FirstImpressionItem:
    lang = (lang or "en").lower()
    usable = 0
    forward = 0
    yaw_vals: List[float] = []

    for pose in seq:
        nose = _fi_get(pose, "NOSE")
        le = _fi_get(pose, "LEFT_EAR")
        re = _fi_get(pose, "RIGHT_EAR")
        if not (_fi_vis_ok(nose) and _fi_vis_ok(le) and _fi_vis_ok(re)):
            continue

        usable += 1
        nx = float(nose[0]); lex = float(le[0]); rex = float(re[0])
        cx = (lex + rex) / 2.0
        span = abs(lex - rex) + 1e-6
        yaw = abs(nx - cx) / span
        yaw_vals.append(yaw)
        if yaw <= 0.18:
            forward += 1

    if usable == 0:
        return FirstImpressionItem(
            title_en="Eye Contact",
            title_th="การสบตา (Eye Contact)",
            scale="—",
            score=0,
            insight_en="Eye contact could not be reliably detected from this clip.",
            insight_th="ระบบไม่สามารถตรวจจับทิศทางการสบตาได้อย่างน่าเชื่อถือจากคลิปนี้",
            impact_en="When the system cannot detect gaze direction, we avoid over-interpreting presence or confidence.",
            impact_th="เมื่อระบบตรวจจับทิศทางการมองไม่ชัดเจน เราจะหลีกเลี่ยงการสรุปเรื่องความมั่นใจหรือความน่าเชื่อถือมากเกินไป",
            details={"usable_frames": 0},
        )

    forward_ratio = forward / usable
    score = int(round(_fi_clamp01((forward_ratio - 0.15) / 0.85) * 100))
    level = _fi_level_from_score(score, lang)

    if score >= 75:
        bullets_en = [
            "Your eye contact is steady, warm, and audience-focused.",
            "You maintain direct gaze during key message points, which increases trust and clarity.",
            "When you shift your gaze, it is done purposefully (e.g., thinking, emphasizing).",
            "There is no sign of avoidance — overall, the eye contact supports confidence and credibility.",
        ]
        bullets_th = [
            "คุณสบตาได้ค่อนข้างสม่ำเสมอ ดูเป็นมิตร และโฟกัสผู้ฟัง",
            "คุณคงการมองตรงในจุดสำคัญ ทำให้เกิดความไว้วางใจและความชัดเจน",
            "เมื่อมีการเปลี่ยนสายตา ดูเป็นการเปลี่ยนอย่างมีเจตนา (เช่น คิด/เน้นประเด็น)",
            "ไม่พบสัญญาณการหลบสายตา — โดยรวมช่วยเสริมความมั่นใจและความน่าเชื่อถือ",
        ]
        impact_en = "Strong eye contact signals presence, sincerity, and leadership confidence, making your message feel more reliable."
        impact_th = "การสบตาที่ดีสะท้อนถึงความอยู่กับปัจจุบัน ความจริงใจ และความมั่นใจแบบผู้นำ ทำให้สารที่สื่อดูน่าเชื่อถือขึ้น"
    elif score >= 45:
        bullets_en = [
            "Your eye contact is generally present, with occasional drift away from the audience.",
            "You return to a forward-facing orientation, but consistency varies across the clip.",
            "A steadier forward gaze would strengthen credibility.",
        ]
        bullets_th = [
            "การสบตาโดยรวมถือว่ามีอยู่ แต่มีบางช่วงที่หลุดจากผู้ฟัง",
            "คุณกลับมามองด้านหน้าได้ แต่ความสม่ำเสมอแตกต่างกันไป",
            "หากคุมการมองด้านหน้าให้สม่ำเสมอขึ้น จะช่วยเพิ่มความน่าเชื่อถือ",
        ]
        impact_en = "More consistent eye contact helps the audience feel included and increases perceived confidence."
        impact_th = "การสบตาที่สม่ำเสมอช่วยให้ผู้ฟังรู้สึกมีส่วนร่วม และเพิ่มการรับรู้เรื่องความมั่นใจ"
    else:
        bullets_en = [
            "Eye contact appears inconsistent, with frequent orientation away from the audience.",
            "This may reduce perceived clarity or create a sense of hesitation.",
            "Stabilizing forward gaze at key message points can significantly improve first impression.",
        ]
        bullets_th = [
            "การสบตาดูไม่สม่ำเสมอ และมีการหันหน้าออกจากผู้ฟังบ่อย",
            "อาจทำให้ความชัดเจนลดลง หรือทำให้ผู้ฟังรู้สึกถึงความลังเล",
            "การคุมการมองด้านหน้าในจังหวะสำคัญจะช่วยยกระดับ First impression ได้มาก",
        ]
        impact_en = "Limited eye contact can weaken trust signals and make the message feel less anchored."
        impact_th = "การสบตาที่จำกัดอาจลดสัญญาณความน่าเชื่อถือ และทำให้สารที่สื่อดูไม่มั่นคง"

    return FirstImpressionItem(
        title_en="Eye Contact",
        title_th="การสบตา (Eye Contact)",
        scale=level,
        score=score,
        insight_en="\n".join(bullets_en),
        insight_th="\n".join(bullets_th),
        impact_en=impact_en,
        impact_th=impact_th,
        details={
            "usable_frames": usable,
            "forward_ratio": float(forward_ratio),
            "yaw_mean": _fi_safe_mean(yaw_vals),
        },
    )


def _analyze_uprightness(seq: List[Dict[str, Tuple[float, float, float]]], lang: str) -> FirstImpressionItem:
    lang = (lang or "en").lower()
    usable = 0
    vert_scores: List[float] = []
    shoulder_tilts: List[float] = []
    head_offsets: List[float] = []

    for pose in seq:
        ls = _fi_get(pose, "LEFT_SHOULDER")
        rs = _fi_get(pose, "RIGHT_SHOULDER")
        lh = _fi_get(pose, "LEFT_HIP")
        rh = _fi_get(pose, "RIGHT_HIP")
        nose = _fi_get(pose, "NOSE")

        if not (_fi_vis_ok(ls) and _fi_vis_ok(rs) and _fi_vis_ok(lh) and _fi_vis_ok(rh) and _fi_vis_ok(nose)):
            continue

        usable += 1
        sx = (float(ls[0]) + float(rs[0])) / 2.0
        sy = (float(ls[1]) + float(rs[1])) / 2.0
        hx = (float(lh[0]) + float(rh[0])) / 2.0
        hy = (float(lh[1]) + float(rh[1])) / 2.0

        dx = sx - hx
        dy = sy - hy
        vert = 1.0 - _fi_clamp01(abs(dx) / (abs(dy) + 1e-6))
        vert_scores.append(vert)

        tilt = abs(float(ls[1]) - float(rs[1]))
        shoulder_tilts.append(tilt)

        head_off = abs(float(nose[0]) - sx)
        head_offsets.append(head_off)

    if usable == 0:
        return FirstImpressionItem(
            title_en="Uprightness (Posture & Upper-Body Alignment)",
            title_th="ความตั้งตรง (ท่าทางและแนวลำตัวส่วนบน)",
            scale="—",
            score=0,
            insight_en="Uprightness could not be reliably detected from this clip.",
            insight_th="ระบบไม่สามารถประเมินแนวลำตัวและความตั้งตรงได้อย่างน่าเชื่อถือจากคลิปนี้",
            impact_en="When posture detection is uncertain, we avoid making strong claims about confidence or authority.",
            impact_th="เมื่อการตรวจจับท่าทางไม่ชัดเจน เราจะหลีกเลี่ยงการสรุปเรื่องความมั่นใจหรือความเป็นผู้นำมากเกินไป",
            details={"usable_frames": 0},
        )

    vert_mean = float(np.mean(vert_scores)) if np is not None else 0.0
    tilt_mean = float(np.mean(shoulder_tilts)) if np is not None else 1.0
    head_mean = float(np.mean(head_offsets)) if np is not None else 1.0

    tilt_score = 1.0 - _fi_clamp01((tilt_mean - 0.01) / 0.06)
    head_score = 1.0 - _fi_clamp01((head_mean - 0.02) / 0.08)
    combined = (0.55 * vert_mean) + (0.25 * tilt_score) + (0.20 * head_score)

    score = int(round(_fi_clamp01(combined) * 100))
    level = _fi_level_from_score(score, lang)

    if score >= 75:
        bullets_en = [
            "You maintain a naturally upright posture throughout the clip.",
            "The chest stays open, shoulders relaxed, and head aligned — signaling balance, readiness, and authority.",
            "Even when you gesture, your vertical alignment remains stable, showing good core control.",
            "There is no visible slouching or collapsing, which supports a professional appearance.",
        ]
        bullets_th = [
            "คุณรักษาท่าทางที่ค่อนข้างตั้งตรงได้ดีตลอดคลิป",
            "อกเปิด ไหล่ผ่อนคลาย และศีรษะอยู่ในแนว — สื่อถึงความสมดุล ความพร้อม และความน่าเชื่อถือ",
            "แม้มีการใช้มือ/gesture แนวตั้งของลำตัวส่วนบนยังคงเสถียร แสดงการคุม core ได้ดี",
            "ไม่พบการหลังค่อมหรือยุบตัวชัดเจน ช่วยให้ภาพรวมดูเป็นมืออาชีพ",
        ]
        impact_en = "Uprightness communicates self-assurance, clarity of thought, and emotional stability—all traits of high-trust communicators."
        impact_th = "ความตั้งตรงของท่าทางสื่อถึงความมั่นใจ ความคิดชัดเจน และความมั่นคงทางอารมณ์ ซึ่งเป็นคุณลักษณะของผู้สื่อสารที่น่าเชื่อถือ"
    elif score >= 45:
        bullets_en = [
            "Your posture is generally upright, with occasional moments of alignment drift.",
            "Shoulders and head remain mostly stable, but consistency varies.",
            "A slightly stronger vertical stack (head–shoulders–hips) would increase authority.",
        ]
        bullets_th = [
            "ท่าทางโดยรวมค่อนข้างตั้งตรง แต่มีบางช่วงที่แนวลำตัวคลาดเคลื่อนเล็กน้อย",
            "ไหล่และศีรษะค่อนข้างเสถียร แต่ความสม่ำเสมอแตกต่างกันไป",
            "หากคุมแนว Head–Shoulders–Hips ให้ตั้งตรงสม่ำเสมอขึ้น จะช่วยเพิ่มภาพลักษณ์ความเป็นผู้นำ",
        ]
        impact_en = "More consistent alignment can elevate presence and reduce perceived uncertainty."
        impact_th = "การจัดแนวลำตัวให้สม่ำเสมอขึ้น จะยกระดับ presence และลดความรู้สึกไม่มั่นใจที่ผู้ฟังอาจรับรู้"
    else:
        bullets_en = [
            "Posture appears less stable, with noticeable alignment changes.",
            "This can read as reduced readiness or lower confidence, even if the content is strong.",
            "Stabilizing the vertical stack and opening the chest can quickly improve first impression.",
        ]
        bullets_th = [
            "ท่าทางดูไม่ค่อยเสถียร และมีการเปลี่ยนแนวลำตัวค่อนข้างชัด",
            "อาจทำให้ดูเหมือนไม่พร้อมหรือมั่นใจลดลง แม้เนื้อหาที่พูดจะดี",
            "การคุมแนวตั้งของลำตัวและเปิดอกจะช่วยปรับ First impression ได้เร็ว",
        ]
        impact_en = "Unstable posture can weaken authority signals and distract from the message."
        impact_th = "ท่าทางที่ไม่เสถียรอาจลดสัญญาณความเป็นผู้นำ และทำให้ผู้ฟังเสียสมาธิจากเนื้อหา"

    return FirstImpressionItem(
        title_en="Uprightness (Posture & Upper-Body Alignment)",
        title_th="ความตั้งตรง (ท่าทางและแนวลำตัวส่วนบน)",
        scale=level,
        score=score,
        insight_en="\n".join(bullets_en),
        insight_th="\n".join(bullets_th),
        impact_en=impact_en,
        impact_th=impact_th,
        details={
            "usable_frames": usable,
            "vertical_mean": vert_mean,
            "tilt_mean": tilt_mean,
            "head_offset_mean": head_mean,
        },
    )


def _analyze_stance(seq: List[Dict[str, Tuple[float, float, float]]], lang: str) -> FirstImpressionItem:
    lang = (lang or "en").lower()
    usable = 0
    base_widths: List[float] = []
    hip_xs: List[float] = []

    for pose in seq:
        la = _fi_get(pose, "LEFT_ANKLE")
        ra = _fi_get(pose, "RIGHT_ANKLE")
        lh = _fi_get(pose, "LEFT_HIP")
        rh = _fi_get(pose, "RIGHT_HIP")
        ls = _fi_get(pose, "LEFT_SHOULDER")
        rs = _fi_get(pose, "RIGHT_SHOULDER")

        if not (_fi_vis_ok(la) and _fi_vis_ok(ra) and _fi_vis_ok(lh) and _fi_vis_ok(rh) and _fi_vis_ok(ls) and _fi_vis_ok(rs)):
            continue

        usable += 1
        ankle_dist = abs(float(la[0]) - float(ra[0]))
        shoulder_width = abs(float(ls[0]) - float(rs[0])) + 1e-6
        bw = ankle_dist / shoulder_width
        base_widths.append(bw)

        hip_x = (float(lh[0]) + float(rh[0])) / 2.0
        hip_xs.append(hip_x)

    if usable == 0:
        return FirstImpressionItem(
            title_en="Stance (Lower-Body Stability & Grounding)",
            title_th="ความมั่นคงของการยืน (ช่วงล่างและการยืนให้ grounded)",
            scale="—",
            score=0,
            insight_en="Stance could not be reliably detected from this clip.",
            insight_th="ระบบไม่สามารถประเมินความมั่นคงของช่วงล่างและการยืนได้อย่างน่าเชื่อถือจากคลิปนี้",
            impact_en="When lower-body detection is uncertain, we avoid over-claiming grounding or stability.",
            impact_th="เมื่อการตรวจจับช่วงล่างไม่ชัดเจน เราจะหลีกเลี่ยงการสรุปเรื่องความมั่นคงมากเกินไป",
            details={"usable_frames": 0},
        )

    bw_mean = float(np.mean(base_widths)) if np is not None else 0.0
    bw_std = float(np.std(base_widths)) if np is not None else 1.0
    sway_std = float(np.std(hip_xs)) if np is not None else 1.0

    bw_center = 1.15
    bw_score = 1.0 - _fi_clamp01(abs(bw_mean - bw_center) / 0.7)
    bw_consistency = 1.0 - _fi_clamp01((bw_std - 0.03) / 0.12)
    sway_score = 1.0 - _fi_clamp01((sway_std - 0.01) / 0.06)

    combined = (0.40 * bw_score) + (0.35 * bw_consistency) + (0.25 * sway_score)
    score = int(round(_fi_clamp01(combined) * 100))
    level = _fi_level_from_score(score, lang)

    if score >= 75:
        bullets_en = [
            "Your stance is symmetrical and grounded, with feet placed about shoulder-width apart.",
            "Weight shifts are controlled and minimal, preventing distraction and showing confidence.",
            "You maintain good forward orientation toward the audience, reinforcing clarity and engagement.",
            "The stance conveys both stability and a welcoming presence, suitable for instructional or coaching communication.",
        ]
        bullets_th = [
            "ท่ายืนของคุณดูสมดุลและ grounded โดยระยะเท้าใกล้เคียงช่วงไหล่",
            "การถ่ายน้ำหนักมีการคุม ไม่แกว่งมาก ลดสิ่งรบกวนและช่วยเสริมความมั่นใจ",
            "ทิศทางลำตัวมุ่งสู่ผู้ฟังได้ดี ช่วยเสริมความชัดเจนและการมีส่วนร่วม",
            "ท่ายืนสื่อถึงความมั่นคงและความเป็นมิตร เหมาะกับการสื่อสารแนวสอน/โค้ช",
        ]
        impact_en = "A grounded stance enhances authority, control, and smooth message delivery, making the speaker appear more prepared and credible."
        impact_th = "การยืนที่มั่นคงช่วยเสริมความเป็นผู้นำ การควบคุมสถานการณ์ และความลื่นไหลในการสื่อสาร ทำให้ดูพร้อมและน่าเชื่อถือ"
    elif score >= 45:
        bullets_en = [
            "Your stance is generally stable, with some variability in base width or subtle shifts.",
            "The overall grounding is present, but consistency could be stronger.",
            "A steadier base can improve perceived confidence.",
        ]
        bullets_th = [
            "ท่ายืนโดยรวมค่อนข้างมั่นคง แต่ยังมีการขยับ/แปรผันบางช่วง",
            "ภาพรวมยัง grounded แต่ความสม่ำเสมอสามารถเพิ่มได้อีก",
            "หากคุมฐานให้เสถียรขึ้น จะช่วยให้ดูมั่นใจมากขึ้น",
        ]
        impact_en = "More consistent grounding reduces distraction and supports clearer communication."
        impact_th = "ความมั่นคงที่สม่ำเสมอช่วยลดสิ่งรบกวนสายตา และทำให้การสื่อสารชัดเจนขึ้น"
    else:
        bullets_en = [
            "Stance appears less grounded, with noticeable variability or sway.",
            "This can distract the audience and reduce perceived confidence.",
            "Stabilizing the base and minimizing shifting during key points can quickly improve first impression.",
        ]
        bullets_th = [
            "ท่ายืนดูไม่ค่อย grounded และมีความแปรผัน/แกว่งชัดเจน",
            "อาจทำให้ผู้ฟังเสียสมาธิและรับรู้ความมั่นใจลดลง",
            "การคุมฐานให้มั่นคงและลดการขยับในจังหวะสำคัญจะช่วยยกระดับ First impression ได้เร็ว",
        ]
        impact_en = "Unstable stance can weaken authority signals and interrupt message flow."
        impact_th = "ท่ายืนที่ไม่มั่นคงอาจลดสัญญาณความเป็นผู้นำ และทำให้จังหวะการสื่อสารสะดุด"

    return FirstImpressionItem(
        title_en="Stance (Lower-Body Stability & Grounding)",
        title_th="ความมั่นคงของการยืน (ช่วงล่างและการยืนให้ grounded)",
        scale=level,
        score=score,
        insight_en="\n".join(bullets_en),
        insight_th="\n".join(bullets_th),
        impact_en=impact_en,
        impact_th=impact_th,
        details={
            "usable_frames": usable,
            "base_width_mean": bw_mean,
            "base_width_std": bw_std,
            "hip_sway_std": sway_std,
        },
    )


def analyze_first_impression(video_path: str, lang: str) -> Tuple[Optional[FirstImpressionResult], Dict[str, Any]]:
    """
    Try MoveNet TFLite first. If it fails, fallback to MediaPipe Pose if available.
    """
    lang = (lang or "en").lower()
    debug: Dict[str, Any] = {
        "enabled": True,
        "primary": "movenet_tflite",
        "fallback": "mediapipe_pose",
        "movenet_model_path": MOVENET_MODEL_PATH,
        "movenet_model_exists": bool(os.path.exists(MOVENET_MODEL_PATH)),
        "mediapipe_available": bool(MP_HAS_SOLUTIONS),
    }

    if cv2 is None or np is None:
        debug["enabled"] = False
        debug["reason"] = "missing_libs"
        return None, debug

    # Primary: MoveNet
    try:
        seq, meta = _fi_extract_pose_sequence_movenet(
            video_path,
            model_path=MOVENET_MODEL_PATH,
            max_frames=1000,
            sample_fps=POSE_SAMPLE_FPS,
        )
        debug["meta"] = meta
        debug["engine_used"] = meta.get("engine", "movenet_tflite")
        debug["frames_with_pose"] = sum(1 for p in seq if p)
    except Exception as e:
        debug["primary_failed"] = True
        debug["primary_error"] = f"{e}"
        log.warning(f"First Impression primary (MoveNet) failed: {e}. Trying MediaPipe fallback...")

        try:
            seq, meta = _fi_extract_pose_sequence_mediapipe(
                video_path,
                max_frames=1000,
                sample_fps=POSE_SAMPLE_FPS,
            )
            debug["meta"] = meta
            debug["engine_used"] = meta.get("engine", "mediapipe_pose")
            debug["frames_with_pose"] = sum(1 for p in seq if p)
        except Exception as e2:
            debug["enabled"] = False
            debug["reason"] = f"both_failed: movenet_error={e}; mediapipe_error={e2}"
            return None, debug

    if not seq or int(debug.get("frames_with_pose", 0) or 0) < 5:
        debug["enabled"] = False
        debug["reason"] = "insufficient_pose_frames"
        return None, debug

    eye = _analyze_eye_contact(seq, lang)
    up = _analyze_uprightness(seq, lang)
    stn = _analyze_stance(seq, lang)

    processed = int(debug.get("meta", {}).get("processed_frames", 0) or 0)
    frames_pose = int(debug.get("frames_with_pose") or 0)
    pose_ratio = frames_pose / max(1, processed)

    note_en = ""
    note_th = ""
    if pose_ratio < 0.35:
        note_en = "Note: Pose visibility was limited in this clip, so First Impression results should be interpreted with caution."
        note_th = "หมายเหตุ: การมองเห็นสเกเลตัน/โพสในคลิปมีข้อจำกัด จึงควรตีความผล First Impression อย่างระมัดระวัง"
    elif pose_ratio < 0.60:
        note_en = "Note: Pose visibility was moderate, so First Impression results may vary depending on camera angle and lighting."
        note_th = "หมายเหตุ: การมองเห็นโพสอยู่ในระดับปานกลาง ผล First Impression อาจแปรผันตามมุมกล้องและแสง"

    return FirstImpressionResult(
        eye_contact=eye,
        uprightness=up,
        stance=stn,
        reliability_note_en=note_en,
        reliability_note_th=note_th,
    ), debug
# =========================
# Main analysis (fallback placeholder)
# =========================
def analyze_video_fallback(video_path: str) -> dict:
    duration = get_video_duration_seconds(video_path)
    total_indicators = int(max(900, min(20000, duration * 30))) if duration else 900

    size = os.path.getsize(video_path) if os.path.exists(video_path) else 12345
    base = (size % 1000) / 1000.0

    def score_from_ratio(r: float) -> int:
        if np is None:
            s = int(round(r * 10))
            return max(1, min(10, s))
        return int(np.clip(round(r * 10), 1, 10))

    engaging_r = 0.47 + (base * 0.10)
    convince_r = 0.52 + (base * 0.12)
    authority_r = 0.49 + (base * 0.10)

    engaging_pos = int(total_indicators * engaging_r)
    convince_pos = int(total_indicators * convince_r)
    authority_pos = int(total_indicators * authority_r)

    engaging_score = score_from_ratio(engaging_pos / max(1, total_indicators))
    convince_score = score_from_ratio(convince_pos / max(1, total_indicators))
    authority_score = score_from_ratio(authority_pos / max(1, total_indicators))
    overall_score = int(round((engaging_score + convince_score + authority_score) / 3.0))

    effort_detection = {"Gliding": 5, "Punching": 3, "Dabbing": 2, "Flicking": 1, "Pressing": 2}
    shape_detection = {"Advancing": 3, "Retreating": 1, "Enclosing": 1, "Spreading": 3, "Directing": 2, "Indirecting": 1}

    for k in list(effort_detection.keys()):
        effort_detection[f"{k.lower()}_count"] = effort_detection[k]
    for k in list(shape_detection.keys()):
        shape_detection[f"{k.lower()}_count"] = shape_detection[k]

    return {
        "duration_seconds": duration,
        "total_indicators": total_indicators,
        "engaging_pos": engaging_pos,
        "convince_pos": convince_pos,
        "authority_pos": authority_pos,
        "engaging_score": engaging_score,
        "convince_score": convince_score,
        "authority_score": authority_score,
        "overall_score": overall_score,
        "effort_detection": effort_detection,
        "shape_detection": shape_detection,
        "analysis_engine": "fallback",
    }


# =========================
# DOCX formatting helpers (match sample)
# =========================
def _docx_set_base_font(doc, lang: str):
    if Pt is None:
        return
    style = doc.styles["Normal"]
    style.font.name = "TH Sarabun New" if (lang or "").startswith("th") else "Calibri"
    style.font.size = Pt(11)


def _docx_apply_para(
    p,
    left: float = 0.0,
    first_line: float = 0.0,
    before: int = 0,
    after: int = 0,
    line: float = 1.15
):
    if Pt is None or Inches is None:
        return
    pf = p.paragraph_format
    pf.left_indent = Inches(left)
    pf.first_line_indent = Inches(first_line)
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def _docx_add_heading(doc, text: str, size: int = 12, bold: bool = True, before: int = 0, after: int = 6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    if Pt is not None:
        r.font.size = Pt(size)
    _docx_apply_para(p, left=0.0, first_line=0.0, before=before, after=after)
    return p


def _docx_add_numbered_line(doc, number: str, text: str):
    p = doc.add_paragraph()
    _docx_apply_para(p, left=0.35, first_line=0.0, before=0, after=4)
    p.add_run(f"{number}\t{text}")
    try:
        p.paragraph_format.tab_stops.add_tab_stop(Inches(0.75))
    except Exception:
        pass
    return p


def _docx_add_subheading(doc, text: str):
    p = doc.add_paragraph()
    _docx_apply_para(p, left=0.55, first_line=0.0, before=8, after=2)
    p.add_run(text)
    return p


def _docx_add_bullet(doc, text: str):
    p = doc.add_paragraph()
    _docx_apply_para(p, left=0.85, first_line=-0.20, before=0, after=6)
    p.add_run("•  " + text)
    return p


def _docx_add_impact_block(doc, label: str, text: str):
    p1 = doc.add_paragraph()
    _docx_apply_para(p1, left=0.95, first_line=0.0, before=10, after=0)
    p1.add_run(label)

    p2 = doc.add_paragraph()
    _docx_apply_para(p2, left=0.95, first_line=0.0, before=0, after=10)
    p2.add_run(text)
    return p1, p2


def _docx_clear_paragraph(p) -> None:
    try:
        for r in list(p.runs):
            try:
                p._element.remove(r._element)  # type: ignore
            except Exception:
                pass
    except Exception:
        pass
    try:
        p.text = ""
    except Exception:
        pass
