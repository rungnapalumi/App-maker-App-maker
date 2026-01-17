# worker.py  --- สำหรับ service ai-people-reader-worker (Render background worker)

import json
import logging
import os
import shutil
import tempfile
import time
from datetime import datetime, timezone
from typing import Any, Dict, Optional

import boto3
from botocore.exceptions import ClientError

# วิดีโอ (ใช้เป็น placeholder – สามารถแทนที่ด้วยโค้ด Johansson/dots ของ Rung ได้เลย)
try:
    import cv2  # type: ignore
    import numpy as np  # type: ignore
except Exception:  # ถ้า import ไม่ได้ จะใช้โหมด copy ตรง ๆ
    cv2 = None  # type: ignore
    np = None  # type: ignore

# สำหรับสร้างไฟล์รายงาน .docx
from docx import Document  # type: ignore

# -------------------------------------------------------------------------
# Config & logger
# -------------------------------------------------------------------------

AWS_BUCKET = os.getenv("AWS_BUCKET") or os.getenv("S3_BUCKET")
AWS_REGION = os.getenv("AWS_REGION", "ap-southeast-1")
POLL_INTERVAL = int(os.getenv("JOB_POLL_INTERVAL", "10"))

if not AWS_BUCKET:
    raise RuntimeError("Missing AWS_BUCKET (or S3_BUCKET) environment variable")

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] [%(name)s] %(message)s",
)
logger = logging.getLogger("worker")

s3 = boto3.client("s3", region_name=AWS_REGION)

JOBS_PENDING_PREFIX = "jobs/pending/"
JOBS_OUTPUT_PREFIX = "jobs/output/"


# -------------------------------------------------------------------------
# Utils
# -------------------------------------------------------------------------

def utc_now_iso() -> str:
    return datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%S.%f%z")


def s3_get_json(key: str) -> Dict[str, Any]:
    resp = s3.get_object(Bucket=AWS_BUCKET, Key=key)
    data = resp["Body"].read()
    return json.loads(data.decode("utf-8"))


def s3_put_json(key: str, payload: Dict[str, Any]) -> None:
    body = json.dumps(payload, ensure_ascii=False).encode("utf-8")
    s3.put_object(
        Bucket=AWS_BUCKET,
        Key=key,
        Body=body,
        ContentType="application/json",
    )


def s3_download_to_tempfile(key: str, suffix: str) -> str:
    """ดาวน์โหลดไฟล์จาก S3 มาเก็บใน temp แล้วคืน path"""
    fd, tmp_path = tempfile.mkstemp(suffix=suffix)
    os.close(fd)

    with open(tmp_path, "wb") as f:
        s3.download_fileobj(AWS_BUCKET, key, f)

    return tmp_path


def s3_upload_file(src_path: str, key: str, content_type: str) -> None:
    with open(src_path, "rb") as f:
        s3.put_object(
            Bucket=AWS_BUCKET,
            Key=key,
            Body=f,
            ContentType=content_type,
        )


def find_one_pending_job_key() -> Optional[str]:
    """หา job JSON ตัวแรกใน jobs/pending/ ที่ลงท้ายด้วย .json"""
    resp = s3.list_objects_v2(Bucket=AWS_BUCKET, Prefix=JOBS_PENDING_PREFIX)
    contents = resp.get("Contents")
    if not contents:
        return None

    for obj in contents:
        key = obj["Key"]
        if key.endswith(".json"):
            return key

    return None


# -------------------------------------------------------------------------
# Video pipeline (placeholder)
# -------------------------------------------------------------------------

def process_video_with_dots(input_path: str, output_path: str) -> None:
    """
    ตัวอย่าง pipeline ง่าย ๆ:
      - อ่านวิดีโอด้วย OpenCV
      - วาดจุดขาวตรงกลางทุกเฟรม
      - บันทึกเป็น MP4 ใหม่

    *** ตรงนี้ Rung สามารถเปลี่ยนเป็นโค้ด Johansson / skeleton / dots ของจริงได้เลย ***
    """
    if cv2 is None:
        # ถ้าไม่มี OpenCV ให้ copy ไฟล์ตรง ๆ ไปก่อน
        shutil.copyfile(input_path, output_path)
        logger.warning("cv2 not available, copied input video directly to output")
        return

    cap = cv2.VideoCapture(input_path)
    if not cap.isOpened():
        raise RuntimeError("Cannot open input video")

    fps = cap.get(cv2.CAP_PROP_FPS) or 25.0
    width = int(cap.get(cv2.CAP_PROP_FRAME_WIDTH) or 1280)
    height = int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT) or 720)

    fourcc = cv2.VideoWriter_fourcc(*"mp4v")
    out = cv2.VideoWriter(output_path, fourcc, fps, (width, height))

    while True:
        ret, frame = cap.read()
        if not ret:
            break

        # วาดจุดขาวตรงกลาง (placeholder)
        center = (width // 2, height // 2)
        cv2.circle(frame, center, 6, (255, 255, 255), thickness=-1)

        out.write(frame)

    cap.release()
    out.release()


# -------------------------------------------------------------------------
# Presentation Skill Report pipeline (TH + EN)
# -------------------------------------------------------------------------

def run_presentation_report_pipeline(
    *,
    input_video_s3_key: str,
    report_s3_key: str,
    candidate_name: Optional[str] = None,
    user_note: Optional[str] = None,
) -> None:
    """
    สร้างไฟล์ .docx รายงาน Presentation Skill แบบสองภาษา
    (ตอนนี้เนื้อหาเป็น template ยังไม่ได้อ่านข้อมูลจากวิดีโอจริง)

    Rung สามารถแก้ฟังก์ชันนี้เพื่อดึง metric จริง ๆ จาก CSV / ML model ทีหลังได้
    """
    doc = Document()

    # Title
    doc.add_heading("AI People Reader – Presentation Skill Report", level=1)
    if candidate_name:
        doc.add_paragraph(f"Candidate: {candidate_name}")
    if user_note:
        doc.add_paragraph(f"Note from evaluator: {user_note}")

    doc.add_paragraph(f"Video source (S3): {input_video_s3_key}")

    # English section
    doc.add_heading("1. Executive Summary (EN)", level=2)
    doc.add_paragraph(
        "This report is generated by AI People Reader. "
        "It focuses on the candidate's presentation skills, "
        "including clarity, engagement, confidence, and structure."
    )

    doc.add_heading("2. Strengths (EN)", level=2)
    doc.add_paragraph("• Clear overall message\n• Consistent eye contact\n• Natural body movement")

    doc.add_heading("3. Areas for Improvement (EN)", level=2)
    doc.add_paragraph(
        "• Vary vocal tone more to keep the audience engaged.\n"
        "• Add more short pauses between key ideas.\n"
        "• Use gestures to emphasize important points, not on every sentence."
    )

    # Thai section
    doc.add_heading("4. สรุปผลการนำเสนอ (TH)", level=2)
    doc.add_paragraph(
        "รายงานฉบับนี้สร้างขึ้นจากระบบ AI People Reader "
        "เพื่อประเมินทักษะการนำเสนองานของผู้สมัคร "
        "โดยเน้นที่ความชัดเจน การดึงดูดผู้ฟัง ความมั่นใจ "
        "และโครงสร้างการเล่าเรื่องโดยรวม"
    )

    doc.add_heading("5. จุดแข็ง (TH)", level=2)
    doc.add_paragraph("• ถ่ายทอดสารหลักได้ชัดเจน\n• สบตาผู้ฟังได้ดี\n• การเคลื่อนไหวเป็นธรรมชาติ")

    doc.add_heading("6. จุดที่ควรพัฒนา (TH)", level=2)
    doc.add_paragraph(
        "• ควรปรับโทนเสียงให้หลากหลายมากขึ้น เพื่อดึงความสนใจของผู้ฟัง\n"
        "• เว้นจังหวะสั้น ๆ ระหว่างประเด็นสำคัญ เพื่อให้ผู้ฟังมีเวลาคิดตาม\n"
        "• ใช้ท่าทางประกอบเฉพาะจุดที่สำคัญ แทนการใช้ตลอดเวลา"
    )

    # Save to temp & upload
    fd, tmp_path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    try:
        doc.save(tmp_path)
        s3_upload_file(
            src_path=tmp_path,
            key=report_s3_key,
            content_type=(
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document"
            ),
        )
        logger.info("Uploaded report to %s", report_s3_key)
    finally:
        try:
            os.remove(tmp_path)
        except OSError:
            pass


# -------------------------------------------------------------------------
# Main job processor
# -------------------------------------------------------------------------

def process_job(job_key: str) -> None:
    logger.info("Processing job JSON: %s", job_key)
    job = s3_get_json(job_key)

    job_id = job.get("job_id")
    input_key = job.get("input_key")
    output_key = job.get("output_key")
    mode = job.get("mode", "dots")
    user_note = job.get("user_note") or ""
    candidate_name = job.get("candidate_name") or job.get("original_filename")

    if not input_key or not output_key:
        raise ValueError("Job JSON missing 'input_key' or 'output_key'")

    # อัปเดตสถานะเป็น processing
    job["status"] = "processing"
    job["updated_at_utc"] = utc_now_iso()
    s3_put_json(job_key, job)

    # เตรียมที่เก็บ report
    report_s3_key = f"{JOBS_OUTPUT_PREFIX}{job_id}/presentation_report.docx"

    try:
        # 1) ดาวน์โหลดวิดีโอ
        in_path = s3_download_to_tempfile(input_key, suffix=".mp4")
        out_fd, out_path = tempfile.mkstemp(suffix=".mp4")
        os.close(out_fd)

        try:
            # 2) ทำวิดีโอ (Johansson / dots ฯลฯ)
            process_video_with_dots(in_path, out_path)

            # 3) อัปโหลด result.mp4 กลับ S3
            s3_upload_file(out_path, output_key, content_type="video/mp4")

        finally:
            # cleanup
            for p in (in_path, out_path):
                try:
                    os.remove(p)
                except OSError:
                    pass

        # 4) สร้างรายงาน TH/EN (สำหรับทุก mode ในตัวอย่างนี้)
        run_presentation_report_pipeline(
            input_video_s3_key=input_key,
            report_s3_key=report_s3_key,
            candidate_name=candidate_name,
            user_note=user_note,
        )

        # 5) อัปเดตสถานะ job
        job["status"] = "finished"
        job["error"] = None
        job["updated_at_utc"] = utc_now_iso()
        job["report_s3_key"] = report_s3_key

        s3_put_json(job_key, job)
        logger.info("Job %s finished", job_id)

    except Exception as e:
        logger.exception("Unexpected error while processing job %s", job_id)
        job["status"] = "failed"
        job["error"] = str(e)
        job["updated_at_utc"] = utc_now_iso()
        s3_put_json(job_key, job)
        # raise ต่อไม่จำเป็น เพราะเราจะวน loop ต่อ


def main_loop() -> None:
    logger.info("Worker started. Bucket=%s Region=%s", AWS_BUCKET, AWS_REGION)
    while True:
        try:
            job_key = find_one_pending_job_key()
            if not job_key:
                time.sleep(POLL_INTERVAL)
                continue

            process_job(job_key)

        except Exception:
            logger.exception("Unexpected error in main loop")
            time.sleep(POLL_INTERVAL)


if __name__ == "__main__":
    main_loop()
